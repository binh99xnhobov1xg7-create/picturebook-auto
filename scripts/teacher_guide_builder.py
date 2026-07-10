"""Teacher's Guide DOCX 生成器 v2（对齐官方 TG_SOP_Level3-4 结构）。

重构（2026-06-05，用户拍板）：按官方 L3-4 SOP 的结构与排版重做，并按级别分档：

  • L0-2（band=low）：保留 Vocabulary Preview + 4 步 Phonics + 全页 Picture Walk，
    但套用新版结构（Lesson Overview 表 / Question Level Guide / Lesson Flow 时间盒 /
    Pause Points + [L0][L1][L2] 分层提问 / GO 模板 / Quality Checklist）。
  • L3-4（band=mid）：完整 SOP（不预教词、按体裁决定 Picture Walk 范围、动态 Pause Points、
    GO 6 模板、写作任务）。
  • L5-6（band=high）：同 mid 结构，认知难度上调。

  ★ 非虚构（non-fiction）可出现在 L3-L6 任意级别：凡 is_nonfiction 为真，
    一律追加 Non-fiction 绘本制作/教学要求板块（参考 L4_B2 Reader：全页 Picture Walk、
    KWL/Sequence GO、科学事实核对），并自动切换 Picture Walk 范围 / Pause Point 标签 / 体裁默认值。

固定模块顺序：
  Lesson Overview → Pre-Reading → During Reading → Post-Reading →
  Portfolio / Extension → Independent Reading → Quality Checklist
  （L3-L6 非虚构追加 Non-fiction Production & Teaching Notes）

100% 英文输出、无中文、无 AI 元评论；标题/指令 tag/核心词加粗；模块间留空行。
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from parser import BookOutline
from text_format import capitalize_names
from tg_quality import (
    assert_behavior_objectives,
    build_tg_context,
    tg_doc_text,
    validate_tg_output,
    validate_tg_preflight,
)


FONT_EN = "Poppins"
FONT_CN = "Alibaba PuHuiTi 2.0 65 Medium"

# ============================================================
#  SOP 配色系统（取自 TG_SOP_Level3-4.docx）
# ============================================================
C_BLUE = "2E74B5"      # 主蓝：大标题 / Part / 表头
C_BLUE_DK = "1F497D"   # 深蓝：小节标题 / [L1]
C_GREEN = "375623"     # 深绿：[L0] / GO 标签
C_RED = "C00000"       # 警示红：[L2] / 决策框头
C_GRAY = "595959"      # 灰：副标题 / caption
C_INK = "1A1A1A"       # 正文近黑

BG_BLUE = "EBF3FB"     # 浅蓝（信息框 / 斑马纹）
BG_BLUE2 = "E8F4FD"    # 浅蓝（词汇 / 原则框）
BG_CREAM = "FFF8E1"    # 奶油（脚本要求 / 提醒框）
BG_PINK = "FCE4EC"     # 浅粉（与红头配）
BG_GREEN = "E2EFDA"    # 浅绿（与绿头配）
BG_GRAY = "F5F5F5"     # 浅灰（strict rule / 斑马纹）
BG_WHITE = "FFFFFF"
BORDER = "D7DEE8"      # 表格细边框（淡蓝灰）

# [L0]/[L1]/[L2] 分层提问 → 颜色（green / blue / red，遵循 SOP）
TAG_COLOR = {"L0": C_GREEN, "L1": C_BLUE_DK, "L2": C_RED}


# ============================================================
#  级别分档 / 体裁
# ============================================================
def _band(level: str) -> str:
    """low = Smart/0/1/2, mid = 3/4, high = 5/6."""
    key = str(level or "").strip().lower()
    if "smart" in key:
        return "low"
    digits = "".join(ch for ch in key if ch.isdigit())
    try:
        n = int(digits)
    except ValueError:
        return "mid"
    if n <= 2:
        return "low"
    if n <= 4:
        return "mid"
    return "high"


_CJK_RE = re.compile(r"[\u4e00-\u9fff\u3000-\u303f\uff00-\uffef]")


def _en(s: str) -> str:
    """从大纲的双语值里抽出纯英文，保证 TG 100% 英文。

    "视觉化 (Visualizing)"            -> "Visualizing"
    "Bubble Map (气泡图)"            -> "Bubble Map"
    "Character & Plot Map (角色情节图)" -> "Character & Plot Map"
    无英文则返回去掉中文后的残余。
    """
    s = (s or "").strip()
    if not s:
        return ""
    m = re.search(r"[（(]([^（）()]+)[)）]", s)
    if m:
        inside = m.group(1).strip()
        outside = (s[: m.start()] + s[m.end():]).strip(" ：:·-，,/")
        cands = [outside, inside]
    else:
        cands = [s]

    def score(c: str) -> tuple:
        letters = len(re.findall(r"[A-Za-z]", c))
        cjk = len(_CJK_RE.findall(c))
        return (letters > 0, -cjk, letters)

    best = max(cands, key=score)
    best = _CJK_RE.sub("", best).strip(" ：:·-，,/")
    best = re.sub(r"\s{2,}", " ", best)
    return best


def _clean_quotes(s: str) -> str:
    """把全角/弯引号统一成直引号，便于英文 TG 排版。"""
    return (s or "").replace("\u201c", '"').replace("\u201d", '"') \
        .replace("\u2018", "'").replace("\u2019", "'") \
        .replace("\u300c", '"').replace("\u300d", '"')


def _syllabus(outline: BookOutline):
    return getattr(outline, "syllabus", None)


def _is_nonfiction(outline: BookOutline) -> bool:
    ft = (getattr(outline, "fiction_type", "") or "").strip().lower()
    if ft:
        return ft.startswith("non")
    rt = (getattr(outline, "reader_type", "") or "").strip().lower()
    return any(x in rt for x in ("non-fiction", "nonfiction", "informational"))


# ============================================================
#  主入口
# ============================================================
def build_teacher_guide(outline: BookOutline, out_path: Path) -> Path:
    band = _band(outline.level)
    is_nf = _is_nonfiction(outline)
    tg_context = build_tg_context(outline, _worksheet_activities(outline))
    validate_tg_preflight(tg_context)

    doc = Document()
    _set_default_font(doc)
    _set_a4_margins(doc)

    # 总标题块（对齐 SOP：大蓝标题 + 书名 + 灰副标 + 蓝分隔线）
    _title_block(doc, outline, is_nf)

    _build_lesson_overview(doc, outline, band, is_nf)
    _build_pre_reading(doc, outline, band)
    if band == "low":
        _build_sor_routine(doc, outline)
    _build_during_reading(doc, outline, band, is_nf)
    _build_post_reading(doc, outline, band, is_nf)
    _build_portfolio(doc, outline, is_nf)
    _build_independent_reading(doc)
    # L3-L6 都可能是虚构/非虚构：非虚构一律追加科普制作与教学要求板块
    if band in ("mid", "high") and is_nf:
        _build_nonfiction_notes(doc, outline)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    validate_tg_output(
        tg_context,
        tg_doc_text(doc),
        _picture_walk_validation_rows(outline, _picture_walk_pages(outline, band, is_nf)),
    )
    doc.save(str(out_path))
    return out_path


# ============================================================
#  Part 1: Lesson Overview
# ============================================================
def _build_lesson_overview(doc, outline: BookOutline, band: str, is_nf: bool) -> None:
    _heading(doc, "Lesson Overview")

    cefr = (outline.cefr or "").strip() or _cefr_default(outline.level)
    lexile = (outline.lexile or "").strip()
    level_cefr = f"Level {_level_label(outline.level)} \u2022 {cefr}"
    if lexile:
        level_cefr += f" \u2022 Lexile {lexile}"
    level_cefr += f" \u2022 Word count: {outline.total_words or '—'}"

    strategy = _reading_strategy(outline, is_nf)
    skill = _reading_skill(outline, is_nf)

    rows = [
        ("Book Title", capitalize_names(outline.title)),
        ("Level / CEFR", level_cefr),
        ("Genre", "Nonfiction" if is_nf else "Fiction"),
        ("Pages", _pages_label(outline)),
        ("Time", "60 minutes"),
        ("Core Vocabulary", _format_vocab(outline)),
        ("Language Focus", _language_focus_label(outline)),
        ("Reading Strategy", strategy),
        ("Reading Skill", skill),
        ("Phonics Focus", _phonics_focus_label(outline)),
    ]
    _kv_table(doc, rows)

    _heading(doc, "Key Objectives", level=3)
    _para(doc, "By the end of this lesson, students will be able to:")
    objectives = _objective_bullets(outline, is_nf)
    assert_behavior_objectives(objectives)
    for obj in objectives:
        _bullet(doc, obj)

    _heading(doc, "Question Level Guide", level=3)
    _callout(
        doc,
        "Comprehension questions in this TG are tagged with one of three levels. "
        "These tags help teachers judge the difficulty of each question and adapt support accordingly.",
        bg=BG_CREAM,
    )
    _grid_table(
        doc,
        ["Tag", "Level", "What the student does", "If student struggles"],
        [
            ["[L0]", "Literal", "Answer is directly visible in the text or picture (What / Where / Who / Point).",
             "Re-read the sentence and point to the word or picture together."],
            ["[L1]", "Inferential", "Combine two pieces of information (How / Why / What happened next).",
             "Ask the two smaller questions that lead to the answer."],
            ["[L2]", "Analytical / Personal Connection", "Reason beyond the text (Evaluate / Compare / Predict / Justify / Connect personally).",
             "Offer a sentence starter and accept any reasoned reply."],
        ],
        widths=[1.6, 3.0, 6.0, 5.4],
    )
    l2_rule = ("Rule: Each Pause Point must include at least one [L0] and one [L1] question when supported by the text. "
               "[L2] questions may appear when pedagogically appropriate. The final Pause Point must include at least one open-ended [L2] question.")
    if band == "high":
        l2_rule += (" At this level, raise the cognitive demand: include [L2] analysis/evaluation at the last "
                    "TWO Pause Points and expect longer, evidence-based responses.")
    elif band == "low":
        l2_rule += " Keep questions mostly literal ([L0]) with light inference; [L2] stays simple and oral."
    _callout(doc, l2_rule, bg=BG_CREAM)

    _heading(doc, "Lesson Flow at a Glance", level=3)
    _grid_table(
        doc,
        ["Stage", "Focus", "Time"],
        _lesson_flow_rows(band),
        widths=[5.0, 9.0, 2.0],
    )


def _lesson_flow_rows(band: str) -> list[list[str]]:
    rows = [
        ["PRE-READING", "", ""],
        ["Warm-up", "Activate prior knowledge and create a reason to read.", "3 min"],
    ]
    if band == "low":
        rows.append(["Phonics + Vocabulary Preview", "Sound work and preview of core words.", "3 min"])
    else:
        rows.append(["Phonics Awareness", "Brief sound recognition check.", "3 min"])
    rows += [
        ["DURING READING", "", ""],
        ["Picture Walk", "Build meaning and introduce vocabulary through pictures.", "5 min"],
        ["Detailed Reading", "Read with comprehension; Pause Points and layered questions.", "15 min"],
        ["Rereading for Fluency", "Build automaticity, expression, and pace.", "5 min"],
        ["POST-READING", "", ""],
        ["Reading Check", "Assess words, fluency, and comprehension.", "10 min"],
        ["Worksheet Practice", "Consolidate vocabulary, patterns, and connections.", "15 min"],
        ["Lesson Close", "Summarize and wrap up with target words.", "4 min"],
    ]
    return rows


# ============================================================
#  Part 2: Pre-Reading
# ============================================================
def _build_pre_reading(doc, outline: BookOutline, band: str) -> None:
    _heading(doc, "PRE-READING")

    _heading(doc, "Warm-up & Purpose \u2022 3 minutes", level=3)
    says, asks, responses = _warmup_script(outline)
    _tagged(doc, "Teacher Says:", says)
    _tagged(doc, "Teacher Asks:", asks)
    _tagged(doc, "Expected Responses:", responses)
    _para_italic(doc, "Do not reveal the story or pre-teach vocabulary out of context here.")

    phonics = outline.phonics or "the target phonics rule from the book"
    if band == "low":
        _heading(doc, "Phonics Focus (Interactive, 4 Steps)", level=3)
        _tagged(doc, "Rule Identification:", phonics)
        _tagged(doc, "Step 1 - Sound Discovery (I Do):", "Teacher models the target sound and segments an example word.")
        _tagged(doc, "Step 2 - Word Family Extension (We Do):", "Practice 2-3 words from the same family together.")
        _tagged(doc, "Step 3 - Text Scavenger Hunt (You Do):", "Students find the target sound on the cover or early pages.")
        _tagged(doc, "Step 4 - Movement Challenge:", "Students clap, tap, or air-write each sound segment.")

        _heading(doc, "Vocabulary Preview", level=3)
        for word in _vocab_words(outline)[:6]:
            _heading(doc, f"{word}:", level=4)
            _tagged(doc, "Teacher Action:", f"Show a clear gesture or picture that represents \"{word}\".")
            _tagged(doc, "Teacher Says:", f"\"This is {word}. Say {word} with me: {word}.\"")
            _tagged(doc, "Expected Response:", f"Students say \"{word}\" and mimic the gesture.")
    else:
        _heading(doc, "Phonics Awareness \u2022 3 minutes", level=3)
        _tagged(doc, "Focus:", _phonics_focus_label(outline) if phonics else phonics)
        _tagged(doc, "Target Words from this book:", _phonics_examples(outline))
        _tagged(doc, "Teacher Says:",
                f"\"Listen carefully. These words share the target sound. Repeat after me: {_phonics_examples(outline)}.\"")
        _tagged(doc, "Quick Task:",
                "\"As we read today, raise your hand quietly when you hear a word with the target sound.\"")
        _callout(doc, "Awareness only - keep this brief. Do not drill. Vocabulary is introduced in context during reading.",
                 bg=BG_CREAM)


# ============================================================
#  Part 2.5: Science of Reading Routine (L0-2 only)
# ============================================================
_SOR_DEFAULT = {
    "before": [
        "Build background knowledge and oral language about the topic.",
        "Preview 2-3 core words with a clear picture or gesture.",
        "Phonological awareness warm-up (1-2 minutes): clap or segment sounds.",
        "Preview the target high-frequency word, if applicable.",
    ],
    "during": [
        "Model decoding behaviours with a short teacher think-aloud.",
        "Finger Tracking: point under each word so students match spoken to printed words.",
        "Echo reading, then choral reading, to build fluency.",
        "Use picture-text integration to support meaning.",
    ],
    "after": [
        "Retell with pictures only.",
        "Sequence the events with visuals.",
        "Reinforce the taught high-frequency word.",
        "Revisit the core vocabulary in context.",
    ],
}


def _sor_lists(outline: BookOutline) -> dict:
    syl = _syllabus(outline)
    strat = getattr(syl, "sor_strategies", None) if syl else None
    out = {}
    for phase in ("before", "during", "after"):
        items = []
        if isinstance(strat, dict):
            for raw in strat.get(phase, []) or []:
                t = _clean_quotes(str(raw)).strip()
                if t and not _CJK_RE.search(t):
                    # 句末补句号，统一观感
                    items.append(t if t[-1] in ".!?" else t + ".")
        out[phase] = items[:5] or list(_SOR_DEFAULT[phase])
    return out


def _build_sor_routine(doc, outline: BookOutline) -> None:
    """L0-2 专属：Science of Reading 前/中/后策略 + 老师指读 + Oral Language Prompts。"""
    _heading(doc, "Science of Reading Routine (Before / During / After)")
    _callout(doc,
             "For emerging readers, run these Science of Reading-aligned moves across the lesson. "
             "Finger Tracking during reading is required: it builds one-to-one print concepts.",
             bg=BG_GREEN, header="\u2605 Emerging-Reader Routine", header_bg=C_GREEN)

    lists = _sor_lists(outline)
    for phase, label in (("before", "Before Reading"), ("during", "During Reading"), ("after", "After Reading")):
        _heading(doc, label, level=3)
        for item in lists[phase]:
            _para(doc, f"\u2022 {item}")

    # Oral Language Prompts（取自大纲；缺失给通用）
    _heading(doc, "Oral Language Prompts", level=3)
    prompts = _oral_prompts(outline)
    for p in prompts:
        _tagged(doc, "Teacher Says:", f"\"{p}\"" if not p.startswith('"') else p)


def _oral_prompts(outline: BookOutline) -> list[str]:
    syl = _syllabus(outline)
    raw = _clean_quotes(getattr(syl, "oral_prompts", "") if syl else "").strip()
    out: list[str] = []
    if raw and not _CJK_RE.search(raw):
        for part in re.split(r"(?<=[.?!])\s+", raw):
            part = part.strip()
            if part:
                out.append(part)
    if not out:
        out = [
            "Point to the picture. What do you see?",
            "Say it with me.",
            "Tell me one thing about this page.",
        ]
    return out[:6]


# ============================================================
#  Part 3: During Reading
# ============================================================
def _build_during_reading(doc, outline: BookOutline, band: str, is_nf: bool) -> None:
    _heading(doc, "DURING READING")

    # --- Picture Walk ---
    _heading(doc, "Step 1: Picture Walk • 5 minutes", level=3)
    _tagged(doc, "Goal:",
            "Build meaning, establish story elements, and introduce vocabulary naturally through pictures before reading the text.")
    walk_pages = _picture_walk_pages(outline, band, is_nf)
    if is_nf or band == "low":
        scope_note = (f"Picture Walk covers Pages {walk_pages[0]}-{walk_pages[-1]} (all pages). "
                      "Why This Arrangement: previewing every page activates background knowledge for each section; "
                      "there is no narrative suspense to protect.")
    else:
        scope_note = _picture_walk_scope_note(outline, walk_pages)
    _callout(doc, scope_note, bg=BG_BLUE,
             header="\U0001F4CC Why this arrangement", header_bg=C_BLUE_DK)
    if band in ("mid", "high"):
        _callout(
            doc,
            "\u2605 SOR Focus - Reading to Learn: At Level 3-4, students move from decoding to comprehension. "
            "Core vocabulary is introduced naturally through pictures and context during reading, never pre-taught.",
            bg=BG_CREAM,
        )

    _picture_walk_table(doc, outline, walk_pages)

    # --- Detailed Reading ---
    _heading(doc, "Step 2: Detailed Reading • 15 minutes", level=3)
    _tagged(doc, "Goal:",
            f"Read the full text and build {_reading_skill(outline, is_nf)} and {_reading_strategy(outline, is_nf)} through structured pause points.")
    _tagged(doc, "Reading Routine:",
            "Teacher reads aloud (model fluency and expression) -> Students echo read -> Students choral read. "
            "Maintain this routine at every page; do not skip any page.")
    if band == "low":
        _tagged(doc, "Finger Tracking:",
                "Point under each word as it is read so students match each spoken word to its printed word "
                "(left-to-right, return sweep). Have students track with their own finger during choral reading.")
    _heading(doc, "Page-by-Page Reading Script", level=4)
    for printed, text in _printed_story_pages(outline):
        _tagged(doc, f"Page {printed}:", f"Teacher reads: \"{text}\"")
    pause_points = _pause_points(outline, is_nf, band)
    for pp in pause_points:
        write_note = (pp.get("write") or "").strip()
        if write_note.lower().startswith("write:"):
            write_note = write_note.split(":", 1)[1].strip()
        lines = [
            f"Read Pages {pp['pages']} aloud. Students echo read, then choral read.",
            ("Why pausing here:", pp["reason"]),
        ]
        for q in pp["questions"]:
            lines.append((f"[{q['tag']}]", q["q"]))
        if pp.get("write"):
            lines.append(("Write:", write_note))
        _callout(doc, lines, bg=BG_WHITE, header=f"\u23f8 {pp['label']} • Pages {pp['pages']}", header_bg=C_BLUE_DK)

    _vocabulary_embedded_check(doc, outline)

    # --- Rereading ---
    _heading(doc, "Step 3: Rereading for Fluency • 5 minutes", level=3)
    if is_nf:
        _tagged(doc, "Round 1 - Expression Reading:",
                "Match tone to information type: a curious voice for questions, a clear confident voice for facts.")
    else:
        _tagged(doc, "Round 1 - Expression Reading:",
                _fluency_emotion_prompt(outline))
    _tagged(doc, "Round 2 - Phrased Reading:",
            "\"Read at a steady pace, in phrases, not word by word. Pause at commas and full stops.\" "
            "Model one sentence, then students read independently or in pairs.")


# ============================================================
#  Part 4: Post-Reading
# ============================================================
def _build_post_reading(doc, outline: BookOutline, band: str, is_nf: bool) -> None:
    _heading(doc, "POST-READING")

    # Reading Check (fixed text)
    _heading(doc, "Reading Check (10 minutes)", level=3)
    _para(doc, "Purpose: Check words understanding, reading fluency, and reading comprehension and expression.")
    _tagged(doc, "Step 1: Words Recognition",
            "Teacher points to the words one by one. Student reads the words. If the student can correctly read "
            "the words, Teacher ticks the words. If the student can't read the words, Teacher helps and puts circles "
            "in the next box.")
    _tagged(doc, "Step 2: Reading Fluency",
            "Student reads the book independently. Teacher listens and marks the words that the student is struggling with.")
    _tagged(doc, "Step 3: Reading Comprehension and Expression",
            "Teacher asks the questions. Student answers the questions independently. Teacher listens and corrects if "
            "necessary. Teacher ticks the questions that the student is able to answer and circles the questions that "
            "the student is struggling with.")

    # Worksheet Practice
    _heading(doc, "WORKSHEET PRACTICE • 15 minutes")
    _para(doc, "Work through each section in order. Teacher provides answers after each section is complete.")
    if band == "low":
        # 5-element analysis per activity (L0-2 SOP)
        for i, ws in enumerate(_worksheet_activities(outline)[:6], 1):
            _heading(doc, f"Activity {i}: {ws.get('title', 'Activity')}", level=4)
            if ws.get("instruction"):
                _tagged(doc, "Goal:", f"Students will {ws['instruction']}")
            _tagged(doc, "Step 1: Modeling (I Do):", "\"Watch me. I will do the first one...\"")
            _tagged(doc, "Step 2: Independent Practice (You Do):", "\"Now you try the rest.\"")
            _tagged(doc, "Expected Student Response:", "Students complete the items using the target language.")
            if ws.get("answer_key"):
                _tagged(doc, "Answer Key:", ws["answer_key"])
    else:
        # Standard sections (L3-6 SOP): Goal + Teacher Briefing + Answer Key
        for i, ws in enumerate(_worksheet_activities(outline)[:6], 1):
            _heading(doc, f"Section {i}: {ws.get('title', 'Activity')}", level=3)
            if ws.get("instruction"):
                _tagged(doc, "Goal:", f"Students will {ws['instruction']}")
            _tagged(doc, "Teacher Briefing:", ws.get("briefing") or
                    "Model the first item aloud, then have students complete the rest independently while you give targeted feedback.")
            if ws.get("answer_key"):
                _callout(doc, f"Answer Key:\n{ws['answer_key']}", bg=BG_CREAM)
            if ws.get("discussion_prompt"):
                _tagged(doc, "Quick Discussion Prompt:", ws["discussion_prompt"])

    # Graphic Organizer guidance (if the worksheet contains a GO-like activity)
    go = _detect_go(outline, is_nf)
    if go:
        _heading(doc, f"Graphic Organizer: {go['name']}", level=4)
        _tagged(doc, "Teacher Briefing:", go["briefing"])
        _tagged(doc, "Word Bank (max 8 words):", go["word_bank"])
        _tagged(doc, "Completion Guidance:", go["guidance"])
        _tagged(doc, "Answer Key:", go["answer_key"])

    # Writing Task guidance (L3-L6 写作类 Worksheet：SOP 完整写作指引)
    if band in ("mid", "high") and _has_writing_task(outline):
        _build_writing_task(doc, outline, band, is_nf)

    # Lesson Close
    _heading(doc, "LESSON CLOSE • 4 minutes")
    _tagged(doc, "Summarize:", _today_objectives(outline, is_nf))
    _tagged(doc, "Reflection Prompt:", _lesson_close_reflection(outline, is_nf))
    _tagged(doc, "Final Wrap-up:",
            f"\"Today we read {capitalize_names(outline.title)} together. Great job using our new words - "
            f"see you next time!\"")


# ============================================================
#  Part 5/6/7 + Non-fiction notes
# ============================================================
def _build_portfolio(doc, outline: BookOutline, is_nf: bool) -> None:
    _heading(doc, "Portfolio / Extension Tasks")
    _para_italic(doc, "Select one option based on available time and student need.")
    if is_nf:
        opt1 = f"Make an illustrated fact poster about {capitalize_names(outline.title)}. Draw 2-3 key facts and label them with the core vocabulary ({_format_vocab(outline)})."
        opt2 = f"Retell {capitalize_names(outline.title)} using the main idea and 2-3 supporting details."
        opt3 = "Write 3-4 true sentences about the topic using at least two core vocabulary words."
    elif _is_l3_plan_story(outline):
        opt1 = "Draw Mia's plan and label each action with a time or order clue: first, on Tuesday, every day."
        opt2 = "Retell the story using Problem - Plan - Result: Mia has many things to do; Mia makes a seven-day plan; Mia feels happy and proud of her plan."
        opt3 = "Write four sentences about your own weekly plan using I will + base verb."
    else:
        opt1 = f"Draw a 3-part comic strip: BEFORE / EVENT / AFTER. Label the feelings in each part using words from the story."
        opt2 = f"Retell {capitalize_names(outline.title)} using the structural framework and at least {_oral_word_count(outline)} core words."
        opt3 = "Write a short paragraph using the book's sentence patterns and at least one core vocabulary word."
    _grid_table(
        doc,
        ["Option", "Task"],
        [
            ["Option 1 / Draw & Label", opt1],
            ["Option 2 / Oral Retell", opt2],
            ["Option 3 / Writing", opt3],
        ],
        widths=[4.0, 12.0],
    )


def _build_independent_reading(doc) -> None:
    _heading(doc, "Independent Reading (Optional)")
    _tagged(doc, "Student Task:", "Choose 2 books from the library.")
    _tagged(doc, "Teacher Prompts:", "\"Look at the pictures.\" \"Try to read.\" \"Tell me one thing about the book.\"")


def _build_nonfiction_notes(doc, outline: BookOutline) -> None:
    """L3-L6 非虚构专属：绘本制作 + 教学要求（参考 L4_B2 Reader 结构）。"""
    _heading(doc, "Non-fiction Production & Teaching Notes")
    _callout(doc, "These notes apply because this is an informational (non-fiction) book.",
             bg=BG_CREAM, header="\u2605 Non-fiction Requirements", header_bg=C_RED)

    _heading(doc, "Scientific Accuracy", level=3)
    _para(doc, "Every fact, number, definition, and illustration must be scientifically accurate and reflect "
               "real-world logic and scale. Verify topic facts against reliable sources before teaching. Pictures "
               "must show realistic proportions (for example, animals and landforms at true relative size).")

    _heading(doc, "Protagonist Explorer Viewpoint", level=3)
    _para(doc, "The series protagonists appear on every page as young science explorers who observe, point to, and "
               "examine the topic. Use this viewpoint to guide students: \"Let's explore what they are looking at.\" "
               "Camera angles vary by content - overhead or bird's-eye views for geography and whole scenes, "
               "close-ups for details.")

    _heading(doc, "Full Picture Walk", level=3)
    _para(doc, "Preview ALL pages during the Picture Walk. Activating background knowledge about every topic "
               "section deepens comprehension; there is no narrative suspense to protect.")

    _heading(doc, "Recommended Graphic Organizers", level=3)
    _para(doc, "Use a KWL chart (Know / Want to know / Learned) to frame inquiry, or a Sequence / Process chart "
               "when the text describes steps or stages. Fill K and W before reading and L after reading.")

    _heading(doc, "Embedded Vocabulary Check", level=3)
    _para(doc, f"After the full text is read, do a brief check on each core word ({_format_vocab(outline)}): ask "
               f"students to find the word in the text and explain what it means here. Accept student-generated meanings.")


def _build_quality_checklist(doc, outline: BookOutline, band: str, is_nf: bool) -> None:
    _heading(doc, "Quality Checklist")
    items = [
        "Output is 100% in English. No other language appears anywhere.",
        "No AI meta-commentary or placeholder text appears anywhere.",
        "Reading Strategy and Reading Skill match the official syllabus names for this book.",
        "Picture Walk scope is stated with a Why This Arrangement explanation.",
        "Pause Point locations are justified with a stated pedagogical reason.",
        "Every Pause Point has at least one [L0] and one [L1] question.",
        "The final Pause Point includes at least one [L2] open-ended question.",
        "Reading Check section is copied exactly from the SOP template.",
        "Independent Reading section is copied exactly from the SOP template.",
        "All core vocabulary words have answer keys in the Worksheet section.",
        "The Graphic Organizer block has Teacher Briefing, Word Bank (max 8), Completion Guidance, and Answer Key.",
        "Portfolio offers exactly 3 options (Creative Arts, Oral, Writing).",
        "Lesson Flow timings add up to 60 minutes.",
        "Bold formatting applied to all tags, headings, and core vocabulary.",
        "Blank lines separate every module for clear visual scanning.",
    ]
    if band == "low":
        items.insert(1, "Vocabulary Preview is included and uses Action / Says / Expected Response.")
        items.insert(2, "Finger Tracking is included in Detailed Reading to build print concepts.")
        items.insert(3, "Science of Reading routine lists Before / During / After moves and Oral Prompts.")
        items.append("Questions stay mostly literal ([L0]) with light inference, matching the level.")
    else:
        items.insert(1, "No Vocabulary Preview section exists; vocabulary is introduced in context.")
    if band == "high":
        items.append("Cognitive demand is raised: extra [L2] analytical/evaluative questions and longer GO responses.")
    if _has_writing_task(outline) and band in ("mid", "high"):
        items.append("Writing Task has a Goal, 3-step Briefing, Sentence Frames with per-blank Word Banks, "
                     "an italic Sample Answer, and a Language Focus note.")
    if is_nf:
        items.append("All facts and illustrations are scientifically accurate and at realistic scale.")
        items.append("Picture Walk previews ALL pages (non-fiction).")
    _grid_table(
        doc,
        ["#", "Check Item", "\u2713 / \u2717"],
        [[str(i + 1), it, ""] for i, it in enumerate(items)],
        widths=[1.0, 13.5, 1.5],
    )


# ============================================================
#  内容生成器
# ============================================================
def _picture_walk_pages(outline: BookOutline, band: str, is_nf: bool) -> list[int]:
    """返回要 Picture Walk 的印刷页码列表（P2 起）。"""
    story = [p for p in outline.pages if p.page_type == "story" and (p.text or "").strip()]
    n = len(story) or 7
    last = n + 1  # 印刷页：story i(1..n) -> Pi+1
    if is_nf or band == "low":
        return list(range(2, last + 1))
    # Fiction L3-L6: preview the setup and turning point, then leave the ending
    # for detailed reading so teachers still have a real comprehension arc.
    preview_end = max(3, last - 2)
    return list(range(2, preview_end + 1))


def _picture_walk_scope_note(outline: BookOutline, walk_pages: list[int]) -> str:
    if _is_l3_plan_story(outline):
        return (
            f"Picture Walk covers Pages {walk_pages[0]}-{walk_pages[-1]}. "
            "Page 7 shows Mia's daily piano practice, and Page 8 reveals her final feeling. "
            "These pages are left for Detailed Reading so students discover the complete plan and emotional resolution through the text."
        )
    return (
        f"Picture Walk covers Pages {walk_pages[0]}-{walk_pages[-1]} only. "
        "Why This Arrangement: later resolution pages are left for Detailed Reading so students discover the ending through the text."
    )


def _story_page(outline: BookOutline, printed_page: int):
    """印刷页码 -> PageSpec（印刷 P2 = story index 1）。"""
    idx = printed_page - 1
    for p in outline.pages:
        if p.index == idx:
            return p
    return None


def _printed_story_pages(outline: BookOutline) -> list[tuple[int, str]]:
    pages: list[tuple[int, str]] = []
    for p in outline.pages:
        if p.page_type == "story" and (p.text or "").strip():
            pages.append((p.index + 1, _clean_quotes((p.text or "").strip())))
    return pages


def _page_vocab_prompt(outline: BookOutline, page_text: str, *, include_sentence: bool = True) -> str:
    text = _clean_quotes(page_text or "").strip()
    words = _vocab_words(outline)
    hits = [
        w for w in words
        if w and re.search(rf"\b{re.escape(w)}\b", text, flags=re.IGNORECASE)
    ]
    if hits:
        word = hits[0]
        if not include_sentence:
            return (
                f"Introduce \"{word}\" briefly from the picture. Students will confirm the exact sentence during Detailed Reading."
            )
        return (
            f"Highlight the word \"{word}\" in context: "
            f"\"Yes, {word} is one of our core words. Listen to the sentence: {text}\""
        )
    if text and include_sentence:
        return f"Confirm the meaning with the actual sentence from this page: \"{text}\""
    if text:
        return "Confirm students' ideas from the picture only. Students will check the exact sentence during Detailed Reading."
    return "Confirm students' ideas from the picture only."


def _pause_points(outline: BookOutline, is_nf: bool, band: str = "mid") -> list[dict]:
    """从 rr_questions 生成 2-3 个 Pause Point，按星级打 [L0]/[L1]/[L2] 标签。

    band=high 时认知上调：除最后一个 Pause 外，倒数第二个 Pause 也补一道 [L2]
    分析/评价题（开放性），体现 L5-6 更高的分析要求。
    """
    if not is_nf and _is_l3_plan_story(outline):
        return _plan_story_pause_points()

    story = [p for p in outline.pages if p.page_type == "story" and (p.text or "").strip()]
    n = len(story) or 7
    last = n + 1

    # 三段页码分组（印刷页）
    cut1 = 2 + max(1, (last - 1) // 3)
    cut2 = 2 + max(2, 2 * (last - 1) // 3)
    groups = [
        (2, min(cut1, last)),
        (min(cut1 + 1, last), min(cut2, last)),
        (min(cut2 + 1, last), last),
    ]
    if is_nf:
        labels = ["Pause 1 - Key Facts 1", "Pause 2 - Key Facts 2", "Pause 3 - Key Facts 3"]
        reasons = [
            "Pause after the first topic section to consolidate the opening facts.",
            "Pause after the second cluster of facts to connect ideas.",
            "Pause at the end to synthesize and extend beyond the text.",
        ]
    else:
        labels = ["Pause 1 - BEFORE", "Pause 2 - EVENT", "Pause 3 - AFTER"]
        reasons = [
            "Pause here because the setting and characters are established.",
            "Pause here because the main problem or event is fully developed.",
            "Pause at the resolution - the primary moment for the Reading Skill focus.",
        ]

    rr = _normalized_rr(outline)
    star_to_tag = {1: "L0", 2: "L1", 3: "L2"}

    pps: list[dict] = []
    for gi, (lo, hi) in enumerate(groups):
        if lo > hi:
            continue
        qs = []
        for q in rr:
            pg = q.get("page")
            if pg is None:
                continue
            if lo <= int(pg) <= hi:
                qs.append({"tag": star_to_tag.get(int(q.get("stars") or 1), "L0"),
                           "q": _qtext(q.get("q", ""))})
        pps.append({
            "label": labels[gi] if gi < len(labels) else f"Pause {gi + 1}",
            "pages": f"{lo}-{hi}" if lo != hi else f"{lo}",
            "reason": reasons[gi] if gi < len(reasons) else "Pause at a natural break in the text.",
            "questions": qs,
            "write": _write_note(gi, is_nf),
        })

    # 把开放性 ⭐⭐⭐（page=None）题放到最后一个 Pause Point，确保 [L2] 出现
    open_q = next((q for q in rr if q.get("page") is None or int(q.get("stars") or 1) >= 3), None)
    if pps:
        last_pp = pps[-1]
        has_l2 = any(qq["tag"] == "L2" for qq in last_pp["questions"])
        if open_q and not has_l2:
            candidate = _qtext(open_q.get("q", ""))
            existing = {
                (qq.get("q") or "").strip().lower()
                for pp in pps for qq in pp.get("questions", [])
            }
            if candidate.strip().lower() not in existing:
                last_pp["questions"].append({"tag": "L2", "q": candidate})
        # 兜底：每个 PP 至少一个 [L0]
        for pp in pps:
            if not any(qq["tag"] == "L0" for qq in pp["questions"]):
                pp["questions"].insert(0, {"tag": "L0", "q": "What do you see on these pages?"})
            if not any(qq["tag"] == "L1" for qq in pp["questions"]):
                pp["questions"].append({"tag": "L1", "q": "Why do you think that happened?"})
        # 高段（L5-6）认知上调：倒数第二个 Pause 也补一道开放性 [L2]
        if band == "high" and len(pps) >= 2:
            penult = pps[-2]
            if not any(qq["tag"] == "L2" for qq in penult["questions"]):
                prompt = ("How does the author's choice here shape the meaning? Use evidence from the text."
                          if is_nf else
                          "What does this moment reveal about the character? Justify with evidence from the text.")
                penult["questions"].append({"tag": "L2", "q": prompt})
        seen_questions: set[str] = set()
        for pp in pps:
            unique = []
            for qq in pp["questions"]:
                key = (qq.get("q") or "").strip().lower()
                if key and key not in seen_questions:
                    seen_questions.add(key)
                    unique.append(qq)
            pp["questions"] = unique
        if pps and not any(qq["tag"] == "L2" for qq in pps[-1]["questions"]):
            pps[-1]["questions"].append({
                "tag": "L2",
                "q": "How does the ending connect to the problem in the story?",
            })
    return [pp for pp in pps if pp["questions"]]


def _write_note(gi: int, is_nf: bool) -> str:
    if is_nf:
        return ["Write: Key facts from the opening section.",
                "Write: New facts and how they connect.",
                "Write: The main idea and one fact that surprised you."][gi if gi < 3 else 2]
    return ["Write: BEFORE - the normal situation and feelings.",
            "Write: EVENT - what happened and the change in feeling.",
            "Write: AFTER - how it ended and the final feeling."][gi if gi < 3 else 2]


def _picture_walk_table(doc, outline: BookOutline, walk_pages: list[int]) -> None:
    rows: list[list[str]] = []
    for printed in walk_pages:
        page = _story_page(outline, printed)
        text = _clean_quotes((page.text if page else "") or "").strip()
        action, ask, expected = _picture_walk_prompts(outline, printed, text)
        vocab_line = _page_vocab_prompt(outline, text, include_sentence=False)
        rows.append([
            f"Page {printed}",
            "\n".join([
                f"Teacher Action: {action}",
                f"Teacher Says: \"Look carefully at this page. What clues can we use before we read?\"",
                f"Teacher Asks: \"{ask}\"",
                f"Expected Response: {expected}",
                f"Teacher Confirms / Expands: {vocab_line}",
                f"Story Element: {_page_story_tag(outline, printed)}",
            ])
        ])
    _grid_table(doc, ["Page", "Picture Walk Script"], rows, widths=[2.4, 13.6])


def _picture_walk_validation_rows(outline: BookOutline, walk_pages: list[int]) -> list[dict]:
    rows = []
    for printed in walk_pages:
        page = _story_page(outline, printed)
        text = _clean_quotes((page.text if page else "") or "").strip()
        action, ask, expected = _picture_walk_prompts(outline, printed, text)
        rows.append({
            "page": printed,
            "teacher_action": action,
            "teacher_asks": ask,
            "expected_response": expected,
        })
    return rows


def _picture_walk_prompts(outline: BookOutline, printed_page: int, text: str) -> tuple[str, str, str]:
    low = (text or "").lower()
    if _is_l3_plan_story(outline):
        if printed_page == 2 or "many things" in low:
            return (
                "Point to Mia and the things around her that suggest a busy week.",
                "What clues show that Mia has many things to do?",
                "Students may notice books, schoolwork, toys, or a busy-looking room.",
            )
        if printed_page == 3 or any(k in low for k in ("homework", "messy")):
            return (
                "Point to the homework and the messy-room clues.",
                "What two jobs can you see Mia may need to finish?",
                "She needs to finish homework and tidy or clean her room.",
            )
        if printed_page == 4 or "show" in low:
            return (
                "Point to the piano or performance clue.",
                "What important event is coming soon? What might Mia need to do before it?",
                "There is a piano show soon, so Mia may need to practice.",
            )
        if printed_page in (5, 6) or any(k in low for k in ("plan", "first", "tuesday")):
            return (
                "Point to the plan/notebook and any order or time clues.",
                "What do you predict Mia will do to organize her week?",
                "Students predict that Mia may make a plan or choose an order for her tasks.",
            )
        if any(k in low for k in ("practice", "hour", "happy", "proud")):
            return (
                "Point to Mia practicing piano and then to her final feeling.",
                "What does Mia do every day? How does the plan help her at the end?",
                "She practices the piano every day. She is happy and proud of her plan.",
            )
    if any(k in low for k in ("worried", "problem", "messy", "lost", "grabbed")):
        return (
            "Point to the character's face and the problem clue in the picture.",
            "What problem can you see? How does the character feel?",
            "Students describe the problem and name a feeling such as worried, sad, or surprised.",
        )
    if any(k in low for k in ("happy", "proud", "help", "solved", "finally")):
        return (
            "Point to the character's final action and facial expression.",
            "What changed from the beginning? How do the characters feel now?",
            "Students notice the solution and describe the final feeling.",
        )
    return (
        "Point to the main character, setting, and one important object.",
        "Who do you see? Where are they? What do you think may happen next?",
        "Students name characters, setting, and a reasonable prediction from the picture.",
    )


def _page_story_tag(outline: BookOutline, printed_page: int) -> str:
    pages = [p for p in outline.pages if p.page_type == "story" and (p.text or "").strip()]
    last_printed = (pages[-1].index + 1) if pages else 8
    if printed_page <= 3:
        return "BEFORE - characters / setting / first feeling"
    if printed_page >= max(2, last_printed - 1):
        return "AFTER - solution / ending feeling"
    return "EVENT - problem / plan / change"


def _is_l3_plan_story(outline: BookOutline) -> bool:
    if _level_label(outline.level) != "3":
        return False
    text = " ".join(_story_sentences(outline)).lower()
    vocab = " ".join(_vocab_words(outline)).lower()
    return (
        "plan" in text
        and any(k in text for k in ("homework", "piano", "practice", "tuesday", "first"))
        and ("homework" in vocab or "practice" in vocab or "plan" in vocab)
    )


def _plan_story_pause_points() -> list[dict]:
    return [
        {
            "label": "Pause 1 - BEFORE",
            "pages": "2-4",
            "reason": "Pause here because Mia's responsibilities and worried feeling are established.",
            "questions": [
                {"tag": "L0", "q": "What things does Mia need to do this week?"},
                {"tag": "L0", "q": "What show is on Sunday?"},
                {"tag": "L1", "q": "Why does Mia feel worried?"},
                {"tag": "L1", "q": "What might help Mia when she has many things to do?"},
            ],
            "write": "BEFORE - Mia has many things to do. She feels worried.",
        },
        {
            "label": "Pause 2 - EVENT",
            "pages": "5-6",
            "reason": "Pause here because Mia starts to solve the problem by making and using a plan.",
            "questions": [
                {"tag": "L0", "q": "What kind of plan does Mia make?"},
                {"tag": "L0", "q": "What will Mia do first?"},
                {"tag": "L1", "q": "How does the plan make Mia's jobs easier?"},
                {"tag": "L0", "q": "Which word tells us what Mia will do before the other tasks?"},
                {"tag": "L0", "q": "Which word tells us the day when Mia will clean her room?"},
            ],
            "write": "EVENT - Mia makes a seven-day plan and starts with homework.",
        },
        {
            "label": "Pause 3 - AFTER",
            "pages": "7-8",
            "reason": "Pause at the resolution because students can connect Mia's plan to her final feeling.",
            "questions": [
                {"tag": "L0", "q": "What will Mia practice every day?"},
                {"tag": "L1", "q": "How does Mia feel at the end?"},
                {"tag": "L1", "q": "How did Mia's feeling change from worried to proud?"},
                {"tag": "L2", "q": "Why is making a plan a good way to solve Mia's problem?"},
            ],
            "write": "AFTER - Mia practices piano every day. She is happy and proud of her plan.",
        },
    ]


def _vocabulary_embedded_check(doc, outline: BookOutline) -> None:
    words = _vocab_words(outline)[:4]
    if not words:
        return
    lines = ["Vocabulary Embedded Check - After Full Reading"]
    for word in words:
        sent = _first_sentence_with_word(outline, word)
        if sent:
            lines.append(f"{word} - Find the sentence: \"{sent}\" What does {word} mean here?")
        else:
            lines.append(f"{word} - Find the word in the book. What does it mean here?")
    _callout(doc, lines, bg=BG_CREAM)


def _first_sentence_with_word(outline: BookOutline, word: str) -> str:
    for s in _story_sentences(outline):
        if re.search(rf"\b{re.escape(word)}\b", s, flags=re.I):
            return s
    return ""


def _normalized_rr(outline: BookOutline) -> list[dict]:
    raw = getattr(outline, "_rr_questions", None) or []
    out = []
    for q in raw:
        if not isinstance(q, dict):
            continue
        out.append({
            "q": str(q.get("q") or q.get("question") or "").strip(),
            "stars": int(q.get("stars") or 1),
            "page": q.get("page"),
        })
    return out


def _qtext(q: str) -> str:
    q = (q or "").strip()
    if not q:
        return "What is happening here?"
    if q[-1] not in "?.!":
        q += "?"
    return capitalize_names(q)


# ---------- Graphic Organizer 检测 ----------
_FEELING_BANK = "happy, excited, worried, nervous, proud, surprised"


def _has_writing_task(outline: BookOutline) -> bool:
    actual_mode = (getattr(outline, "_worksheet_second_reading_mode", "") or "").lower()
    if actual_mode:
        return actual_mode in {"writing", "writing_official"}
    wd = getattr(outline, "_worksheet_data", None)
    for ws in (getattr(outline, "_worksheet_questions", None) or []):
        if isinstance(ws, dict) and "writ" in (ws.get("type") or "").lower():
            return True
    if isinstance(wd, dict) and wd.get("writing") and not actual_mode:
        return True
    return False


def _build_writing_task(doc, outline: BookOutline, band: str, is_nf: bool) -> None:
    """SOP 写作任务完整指引：Goal / 3 步 Briefing / Sentence Frames(每空 Word Bank) /
    斜体 Sample / Language Focus。内容扣住本书的句型与核心词。"""
    wd = getattr(outline, "_worksheet_data", None)
    writing = wd.get("writing") if isinstance(wd, dict) else {}
    min_w = (writing or {}).get("min_words", 50 if band == "mid" else 60)
    max_w = (writing or {}).get("max_words", 80 if band == "mid" else 110)
    vocab = _vocab_words(outline)
    vbank = ", ".join(vocab[:6]) or "the core words from the book"
    pattern = _en(getattr(_syllabus(outline), "sentence_pattern", "") if _syllabus(outline) else "") \
        or outline.grammar_focus

    _heading(doc, "Writing Task", level=4)
    _tagged(doc, "Goal:",
            f"Students will {'write 2-3 facts about the topic' if is_nf else 'retell the story'} in "
            f"{min_w}-{max_w} words, using the book's sentence patterns and core vocabulary "
            f"({vbank}).")

    _tagged(doc, "Teacher Briefing - Step 1 (Set Up):",
            "Read the prompt and the planning box together. Make sure students know what to write before they start.")
    _tagged(doc, "Teacher Briefing - Step 2 (Model, I Do):",
            "Think aloud and write ONE sentence on the board using the first frame, then erase it so students write their own.")
    _tagged(doc, "Teacher Briefing - Step 3 (Independent, You Do):",
            "Students write independently while you confer; prompt with the frames and Word Banks as needed.")

    # Sentence Frames + 每空 Blank Type Label + 每空 Word Bank(<=6)
    if is_nf:
        frames = [
            ("I learned that [fact]. ", "[fact] = one true detail from the book", vbank),
            ("One important fact is [fact], and it [verb]. ", "[verb] = action/process word", vbank),
            ("This matters because [reason].", "[reason] = why it is important", "use your own idea"),
        ]
    else:
        frames = [
            ("At the beginning, [character] felt [feeling] because [reason]. ", "[feeling] = emotion word", _FEELING_BANK),
            ("Then, [event] happened, so [character] [action]. ", "[event] = the main change", vbank),
            ("In the end, [character] felt [feeling] and learned [lesson].", "[lesson] = the message", _FEELING_BANK),
        ]
    _tagged(doc, "Sentence Frames:", f"Provide these frames; each blank has a label and a small Word Bank.{(' Pattern: ' + pattern) if pattern else ''}")
    for frame, blank_label, bank in frames:
        _heading(doc, frame.strip(), level=4)
        _tagged(doc, "Blank Type:", blank_label)
        _tagged(doc, "Word Bank (max 6):", bank)

    if is_nf:
        sample = ("I learned that bats sleep in the day. One important fact is that they hunt at night, and it helps "
                  "control insects. This matters because it keeps nature in balance.")
    else:
        sample = _story_sample_answer(outline)
    _tagged(doc, "Sample Answer:", "")
    _para_italic(doc, sample)

    _tagged(doc, "Language Focus:",
            ("Use clear topic sentences and linking words (because, so, also). Keep facts accurate."
             if is_nf else
             "Use the book's sentence pattern and past-tense verbs; link ideas with because, so, and then."))


def _go_by_name(name: str, word_bank: str, is_nf: bool, base) -> dict | None:
    """按官方 GO 名称关键词匹配六大模板，命中返回四要素指导（保留大纲原名）。"""
    n = (name or "").lower()

    def b(briefing, guidance, answer):
        return base(name, briefing, guidance, answer)

    if "venn" in n or ("compare" in n and "contrast" in n):
        return b("\"The left circle is for things ONLY true about the first item, the right ONLY about the second, "
                 "the middle for BOTH. Let's find one for the middle together, then you try the rest.\"",
                 "Each section needs at least 2 points. Accept single words or short phrases.",
                 "Provide 2-3 points for each section (left, middle, right) based on the book.")
    if any(k in n for k in ("sequence", "timeline", "process", "flow", "cycle", "step")):
        return b("\"This chart shows what happened in order - first, next, then, last. Let's find the FIRST one "
                 "together, then you fill in the rest.\"",
                 "Each box: one short sentence or phrase. Use connectors First / Next / Then / Finally.",
                 "Provide the complete ordered sequence for all boxes from the book.")
    if "plan" in n and "chart" in n:
        return b("\"This chart shows the character's plan. A good plan tells us two things: what to do and when to do it. Let's do the first row together: Mia will do homework first. Now you find the other actions and time clues in the story.\"",
                 "Each row needs one action and one time/order clue. Keep answers short. Students should look back at Pages 6-7 for evidence.",
                 "homework - do homework - first; room - clean her room - on Tuesday; piano - practice the piano - every day; piano - play for one hour - every day.")
    if "kwl" in n:
        return b("\"K is what you already knew. W is what you wanted to find out. L is what you learned from the book. "
                 "Fill K and W before reading and L after reading.\"",
                 "K: 1-2 prior facts. W: 1 question. L: 2-3 key facts from the book using the Word Bank.",
                 "Provide a model L column (K and W are student-generated).")
    if "story element" in n or "story map" in n or ("character" in n and "plot" in n):
        return b("\"Five boxes: Characters, Setting, Problem, Solution, Feelings. Let's do Characters together, "
                 "then you try the rest.\"",
                 "Characters: names only. Setting: where + when. Problem/Solution: one sentence each. "
                 "Feelings: beginning -> middle -> end.",
                 "Provide complete answers for all five boxes from the book.")
    if any(k in n for k in ("bubble", "web", "main idea", "mind map", "concept", "describ", "spider")):
        return b("\"The centre is the main topic. Each outer bubble holds one detail that describes it. "
                 "Let's add the first detail together, then you add the rest.\"",
                 "Centre: the topic in 1-3 words. Each bubble: one describing detail or example from the book (aim for 4-6).",
                 "Provide the centre topic and 4-6 supporting detail bubbles from the book.")
    if any(k in n for k in ("fact", "opinion", "t-chart", "t chart")):
        return b("\"On the left we write FACTS (things we can check). On the right we write OPINIONS (what someone "
                 "thinks or feels). Let's sort the first one together.\"",
                 "Each column needs at least 2 entries, each a short sentence taken or inferred from the text.",
                 "Provide 2-3 facts and 2-3 opinions based on the book.")
    if any(k in n for k in ("tree", "classif", "categor", "sort", "group")):
        return b("\"The top is the big group. Each branch is a smaller category, and under it we list examples. "
                 "Let's label the first branch together.\"",
                 "Each branch: one category label + 2-3 examples drawn from the book.",
                 "Provide the category labels and example members for every branch from the book.")
    if any(k in n for k in ("before", "beginning", "middle", "end", "b-m-e", "bme")):
        return b("\"BEFORE is what was normal at the start, EVENT is the unexpected thing that happened, AFTER is how "
                 "it ended and how people felt. Let's do BEFORE together, then you do EVENT and AFTER.\"",
                 "Each box needs 1-2 sentences. BEFORE: setting + emotion. EVENT: change + emotion. AFTER: resolution + feeling.",
                 "Provide a complete model answer for all three boxes from the book.")
    return None


def _detect_go(outline: BookOutline, is_nf: bool) -> dict | None:
    """GO 指导：优先用官方大纲指定的 GO 名（六模板匹配 / 未列出自生成四要素），
    无大纲 GO 时回退到从 worksheet 题型检测。"""
    vocab = _vocab_words(outline)
    word_bank = ", ".join(vocab[:8]) or "(select up to 8 key words from the book)"

    def base(name, briefing, guidance, answer):
        return {"name": name, "briefing": briefing, "word_bank": word_bank,
                "guidance": guidance, "answer_key": answer}

    actual_mode = (getattr(outline, "_worksheet_second_reading_mode", "") or "").lower()
    if actual_mode == "planchart":
        return {
            "name": "Plan Chart",
            "briefing": "\"This chart shows Mia's plan. A good plan tells us two things: what to do and when to do it. "
                         "Let's do the first row together: Mia will do homework first. Now you find the other actions and time clues in the story.\"",
            "word_bank": "first, clean her room, every day, play for one hour",
            "guidance": "Each row needs one action and one time/order clue. Keep answers short. Students should look back at Pages 6-7 for evidence.",
            "answer_key": "homework - do homework - first; room - clean her room - on Tuesday; piano - practice the piano - every day; "
                          "piano - play for one hour - every day.",
        }
    if actual_mode == "timeline":
        return base(
            "Sequence / Timeline Chart",
            "\"This chart shows what happened in order. Let's find the FIRST one together, then you fill in "
            "the rest.\"",
            "Each box: one short sentence or phrase from the book. Use sequence words only when they fit the text.",
            "Provide the complete ordered sequence from the book.")

    # 1) 大纲精确 GO（命中 syllabus 时）：按名称匹配六模板，未列出走通用四要素
    syl_go = _en(getattr(outline, "graphic_organizer", ""))
    if syl_go:
        tpl = _go_by_name(syl_go, word_bank, is_nf, base)
        if tpl:
            return tpl
        # 未列出的 GO 类型：通用四要素（Teacher Briefing + Word Bank + Completion Guidance + Answer Key）
        return base(
            syl_go,
            f"\"This organizer is called the {syl_go}. Let's complete the first part together, then you finish "
            "the rest using the book.\"",
            "Each section needs 1-2 short points or sentences drawn directly from the text.",
            "Provide a complete model based on the book for every section of the organizer.")

    # 2) 回退：从 worksheet 题型检测
    types = set()
    for ws in (getattr(outline, "_worksheet_questions", None) or []):
        if isinstance(ws, dict):
            types.add((ws.get("type") or "").lower())

    if "compare_contrast" in types:
        return base(
            "Venn Diagram (Compare and Contrast)",
            "\"The left circle is for things ONLY true about A, the right circle ONLY about B, the middle for BOTH. "
            "Let's find one for the middle together, then you try the rest.\"",
            "Each section needs at least 2 points. Accept single words or short phrases.",
            "Provide 2-3 points for each section (left, middle, right) based on the book.")
    if types & {"story_sequence", "word_order", "word_order_simple", "sequence"}:
        return base(
            "Sequence / Timeline Chart",
            "\"This chart shows what happened in order - first, next, then, last. Let's find the FIRST one "
            "together, then you fill in the rest.\"",
            "Each box: one short sentence or phrase. Use connectors First / Next / Then / Finally.",
            "Provide the complete ordered sequence for all boxes from the book.")
    if is_nf:
        return base(
            "KWL Chart (Know / Want to know / Learned)",
            "\"K is what you already knew. W is what you wanted to find out. L is what you learned from the book. "
            "Fill K and W before reading and L after reading.\"",
            "K: 1-2 prior facts. W: 1 question. L: 2-3 key facts from the book using the Word Bank.",
            "Provide a model L column (K and W are student-generated).")
    if types & {"plot_chart", "plot_chart_pbl"}:
        return base(
            "Story Elements Chart",
            "\"Five boxes: Characters, Setting, Problem, Solution, Feelings. Let's do Characters together, "
            "then you try the rest.\"",
            "Characters: names only. Setting: where + when. Problem/Solution: one sentence each. "
            "Feelings: beginning -> middle -> end.",
            "Provide complete answers for all five boxes from the book.")
    return base(
        "Before - Event - After Chart",
        "\"BEFORE is what was normal at the start, EVENT is the unexpected thing that happened, AFTER is how it "
        "ended and how people felt. Let's do BEFORE together, then you do EVENT and AFTER.\"",
        "Each box needs 1-2 sentences. BEFORE: setting + emotion. EVENT: change + emotion. AFTER: resolution + feeling.",
        "Provide a complete model answer for all three boxes from the book.")


# ---------- Worksheet / vocab / objectives 复用 ----------
def _worksheet_activities(outline: BookOutline) -> list[dict]:
    lvl = _level_label(outline.level)
    if lvl == "3":
        return _l3_standard_worksheet_activities(outline)
    qs = getattr(outline, "_worksheet_questions", None) or []
    out: list[dict] = []
    for ws in qs[:6]:
        if not isinstance(ws, dict):
            continue
        ak = ws.get("answer_key")
        out.append({
            "title": capitalize_names(ws.get("title") or _type_title(ws.get("type", "Activity"))),
            "instruction": ws.get("instruction", "") or _type_goal(ws.get("type", "")),
            "extra": ws.get("extra", ""),
            "briefing": "",
            "answer_key": _format_answer_key(ak, ws.get("items")),
        })
    return out


def _l3_standard_worksheet_activities(outline: BookOutline) -> list[dict]:
    words = _vocab_words(outline)[:4]
    if len(words) < 4:
        words = (words + ["week", "homework", "plan", "practice"])[:4]
    story = " ".join(_story_sentences(outline)).lower()
    is_plan = any(k in story for k in ("homework", "seven-day plan", "piano", "on tuesday"))
    return [
        {
            "title": "Vocabulary — Missing Letters",
            "instruction": "look at each picture clue and complete the target word",
            "briefing": "Say each picture clue aloud. Students complete the missing letters, then read the whole word.",
            "answer_key": "; ".join(f"{i + 1}. {w}" for i, w in enumerate(words)),
        },
        {
            "title": "Vocabulary — Fill in the Blanks",
            "instruction": "use the core vocabulary words in story sentences",
            "briefing": "Read each sentence first. Students choose one word from the word bank and write it in the blank.",
            "answer_key": "; ".join(f"{i + 1}. {w}" for i, w in enumerate(words)),
        },
        {
            "title": "Sentences — Complete the Frame",
            "instruction": "complete sentences using the target sentence frame",
            "briefing": "Use the example sentence first. Students choose the correct action phrase for each sentence.",
            "answer_key": "1. do homework; 2. clean her room; 3. practice the piano; 4. play",
        },
        {
            "title": "Sentences — Write Your Own Plan",
            "instruction": "write short personal sentences using the target frame",
            "briefing": "Remind students that answers can be their own ideas. Check that each sentence starts with I will.",
            "answer_key": "Open response. Accept any clear plan sentence that follows the frame I will + action.",
        },
        {
            "title": "Reading — True or False",
            "instruction": "read the story and mark each statement T or F",
            "briefing": "Students look back at the story before choosing T or F. Ask them to point to the sentence that proves the answer.",
            "answer_key": "1. T; 2. F; 3. T; 4. F",
        },
        {
            "title": "Reading — Plan Chart" if is_plan else "Reading — Graphic Organizer",
            "instruction": "complete the graphic organizer using clues from the story",
            "briefing": "Do the first row together. Students complete the remaining blanks using the word bank and story evidence.",
            "word_bank": ["first", "clean her room", "every day", "play for one hour"] if is_plan else [],
            "answer_key": ("1. first; 2. clean her room; 3. every day; 4. play for one hour"
                           if is_plan else "Use short phrases from the story to complete each blank."),
            "discussion_prompt": ("How does Mia's plan help her feel happy and proud at the end?"
                                  if is_plan else "How does this organizer help you understand the book?"),
        },
    ]


def _type_title(t: str) -> str:
    return (t or "Activity").replace("_", " ").title()


def _type_goal(t: str) -> str:
    t = (t or "").lower()
    mapping = {
        "match_definition": "match each word to its meaning",
        "fill_blank": "complete the sentences with the correct word",
        "inference": "answer questions about the story",
        "true_false": "decide whether each statement is true or false",
        "story_sequence": "put the events in order",
        "compare_contrast": "compare and contrast two things from the book",
    }
    return mapping.get(t, "practice the target language from the book")


def _answer_text_for_item(it: dict) -> str:
    """从一条题目 item 解析"对齐后的正确答案文本"（与 worksheet 同源）。

    关键修复：MC/inference 的 correct 是【选项下标】，必须解析成对应的选项文字，
    而不是把下标 0/1/2 直接当答案打印。其余类型按 answer / statement 字段取。
    """
    # 1) 直接给了 answer
    a = it.get("answer")
    if a is not None and str(a).strip():
        return str(a).strip()
    # 2) MC：correct 下标 → 选项文字
    opts = it.get("options")
    if isinstance(opts, list) and opts and it.get("correct") is not None:
        try:
            ci = int(it.get("correct"))
            if 0 <= ci < len(opts):
                return f"{chr(65 + ci)}. {str(opts[ci]).strip()}"
        except (TypeError, ValueError):
            pass
    # 3) 真假题
    if it.get("statement") and it.get("answer"):
        return str(it.get("answer")).strip()
    # 4) 排序题
    if it.get("order") is not None and (it.get("event") or it.get("text")):
        return f"({it.get('order')}) {it.get('event') or it.get('text')}"
    return ""


def _format_answer_key(ak, items) -> str:
    vals: list[str] = []
    if isinstance(ak, list) and ak:
        vals = [str(a) for a in ak if str(a).strip()]
    elif isinstance(ak, dict) and ak:
        vals = [f"{k}: {v}" for k, v in ak.items()]
    elif isinstance(items, list):
        for it in items:
            if isinstance(it, dict):
                v = _answer_text_for_item(it)
                if v:
                    vals.append(v)
    if not vals:
        return ""
    return "; ".join(f"{i + 1}. {v}" for i, v in enumerate(vals))


def _vocab_words(outline: BookOutline) -> list[str]:
    if outline.vocabulary_simple:
        return outline.vocabulary_simple
    if outline.vocabulary_mastery:
        return outline.vocabulary_mastery + outline.vocabulary_exposure
    return []


def _format_vocab(outline: BookOutline) -> str:
    return ", ".join(_vocab_words(outline)) or "the target vocabulary"


def _pages_label(outline: BookOutline) -> str:
    pages = _printed_story_pages(outline)
    if not pages:
        return "Book pages unavailable"
    lo, hi = pages[0][0], pages[-1][0]
    return f"{len(pages) + 1} story pages: Pages {lo}-{hi}, plus title and back cover"


def _language_focus_label(outline: BookOutline) -> str:
    raw = outline.grammar_focus or _en(getattr(_syllabus(outline), "sentence_pattern", "") if _syllabus(outline) else "")
    raw = _clean_quotes(raw).strip()
    if not raw:
        return "Target sentence patterns from the book"
    if "[subject]" in raw.lower() and "will" in raw.lower():
        return "Subject + will + base verb"
    raw = raw.replace("[Subject]", "Subject").replace("[subject]", "subject")
    raw = raw.replace("[action]", "base verb").replace("[Action]", "base verb")
    return raw


def _phonics_focus_label(outline: BookOutline) -> str:
    raw = _clean_quotes(outline.phonics or "").strip()
    if not raw:
        return "Target phonics rule from the book"
    low = raw.lower()
    if "ay" in low:
        return 'Long /eɪ/ sound spelled "ay," as in day'
    if "ee" in low:
        return 'Long /iː/ sound spelled "ee," as in sheep'
    return raw


def _theme_fallback(outline: BookOutline, is_nf: bool) -> str:
    if is_nf:
        return "Informational Reading"
    if _is_l3_plan_story(outline):
        return "Planning and Responsibility"
    skill = _reading_skill(outline, is_nf)
    return skill or "Reading to Learn"


def _warmup_script(outline: BookOutline) -> tuple[str, str, str]:
    title = capitalize_names(outline.title)
    if _is_l3_plan_story(outline):
        return (
            f"\"Today we are going to read {title}. Before we read, let's think about busy days and how people organize what they need to do.\"",
            "\"When you have many things to do in one week, what can help you remember them? Have you ever made a simple plan or list?\"",
            "Students may say: make a list, use a calendar, ask a parent, do one thing first, practice every day.",
        )
    if "predict" in _reading_strategy(outline, _is_nonfiction(outline)).lower():
        return (
            f"\"Today we are going to read {title}. We will use the title and pictures to make careful predictions before we read.\"",
            "\"What can a title or picture help you guess? What should we do after we read to check our guess?\"",
            "Students may say: guess what happens, look for clues, read to check, change the prediction.",
        )
    return (
        f"\"Today we are going to read {title}. Let's connect the topic to something you already know before we read.\"",
        "\"Have you seen or experienced something like this topic before? What detail would you like to find out?\"",
        "Students share brief personal experience or curiosity. Accept all reasonable responses.",
    )


def _fluency_emotion_prompt(outline: BookOutline) -> str:
    if _is_l3_plan_story(outline):
        return (
            "Match voice to the emotions stated or supported by the book: concerned when Mia has many things to do, "
            "worried when she makes the plan, clear and steady for the plan sentences, and happy/proud on the final page."
        )
    return "Match voice to the feelings shown in the text and pictures. Use a different voice only when the book gives evidence for that feeling."


def _story_sentences(outline: BookOutline) -> list[str]:
    sentences: list[str] = []
    for _, page_text in _printed_story_pages(outline):
        for part in re.split(r"(?<=[.!?])\s+", page_text):
            s = _clean_quotes(part).strip()
            if s:
                sentences.append(s)
    return sentences


def _story_sample_answer(outline: BookOutline) -> str:
    sentences = _story_sentences(outline)
    if not sentences:
        return (
            f"At the beginning, the story introduces {capitalize_names(outline.title)}. "
            "Then, the character faces a problem and makes a plan. In the end, the problem is solved."
        )
    first = sentences[0]
    middle = sentences[len(sentences) // 2] if len(sentences) > 2 else sentences[-1]
    last = sentences[-1]
    return f"At the beginning, {first} Then, {middle} In the end, {last}"


def _oral_word_count(outline: BookOutline) -> int:
    return min(4, max(2, len(_vocab_words(outline)) or 4))


def _build_objectives(outline: BookOutline) -> str:
    words = _vocab_words(outline)
    word_str = ", ".join(words[:4]) if words else "the target vocabulary"
    grammar = outline.grammar_focus or "the target sentence frames"
    phonics = outline.phonics or "the target phonics rule"
    return (
        f"Students will be able to identify and use the vocabulary words {word_str}; "
        f"recognize the phonics rule ({phonics}); use the grammar pattern {grammar}; "
        f"answer comprehension questions about the text; and express their own ideas "
        f"using the vocabulary and patterns from the book."
    )


def _objective_bullets(outline: BookOutline, is_nf: bool) -> list[str]:
    words = ", ".join(_vocab_words(outline)[:4]) or "the core vocabulary"
    grammar = _language_focus_label(outline)
    if _is_l3_plan_story(outline):
        return [
            "Identify Mia's problem, plan, and final feeling.",
            "Match four actions with their correct time or order clues.",
            f"Use {words} in story-based sentences.",
            f"Use {grammar} to describe a future plan.",
            "Make and confirm one prediction using evidence from the book.",
        ]
    return [
        f"Identify and use core vocabulary in context: {words}.",
        f"Use the language focus in a sentence: {grammar}.",
        "Answer comprehension questions using evidence from the book.",
        f"Practice {_reading_skill(outline, is_nf)} through pause points and worksheet tasks.",
        "Retell or organize the text using the worksheet graphic organizer.",
    ]


def _bullet(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _font(p.add_run(capitalize_names(text)), size_pt=11, color=C_INK)


def _today_objectives(outline: BookOutline, is_nf: bool) -> str:
    skill = _reading_skill(outline, is_nf)
    strategy = _reading_strategy(outline, is_nf)
    words = ", ".join(_vocab_words(outline)[:4]) or "the target vocabulary"
    if _is_l3_plan_story(outline):
        return (
            "\"Today you identified Mia's problem, plan, and final feeling. You used time and order clues "
            "to understand her seven-day plan, and you checked whether your predictions were correct.\""
        )
    return (f"\"Today students were able to read and understand {capitalize_names(outline.title)}, "
            f"use the words {words}, and practice {skill} and {strategy}.\"")


def _lesson_close_reflection(outline: BookOutline, is_nf: bool) -> str:
    if _is_l3_plan_story(outline):
        return "\"What prediction did you make about Mia's plan? Which page confirmed or changed your idea?\""
    strategy = _reading_strategy(outline, is_nf)
    skill = _reading_skill(outline, is_nf)
    return f"\"How did {strategy} help you understand {skill} in this book? Which page gives your evidence?\""


def _phonics_examples(outline: BookOutline) -> str:
    text = _clean_quotes(outline.phonics or "").strip()
    keys = [
        key for key in ("igh", "ay", "ee", "ai", "oa", "ow", "ou", "oo", "sh", "ch", "th", "ph", "ck")
        if re.search(rf"(?<![A-Za-z]){re.escape(key)}(?![A-Za-z])", text, flags=re.IGNORECASE)
    ]
    if not keys:
        quoted = re.findall(r'"([^"]+)"|\'([^\']+)\'', text)
        keys = [next((part for part in tup if part), "") for tup in quoted]
    keys = [k.lower().strip() for k in keys if k and len(k.strip()) <= 4]

    story_words = []
    seen = set()
    if keys:
        for s in _story_sentences(outline):
            for w in re.findall(r"[A-Za-z][A-Za-z'-]*", s):
                clean_w = w.strip("'")
                if "-" in clean_w:
                    parts = [part for part in clean_w.split("-") if part]
                    clean_w = next((part for part in reversed(parts) if any(key in part.lower() for key in keys)), clean_w)
                lw = clean_w.lower()
                if lw in seen:
                    continue
                if any(key in lw for key in keys):
                    seen.add(lw)
                    story_words.append(clean_w)
    if story_words:
        return ", ".join(story_words[:4])
    m = re.search(r"\(([^)]+)\)", text)
    if m:
        return m.group(1)
    return "2-3 words from the book"


def _reading_strategy(outline: BookOutline, is_nf: bool) -> str:
    """优先用官方大纲精确策略名（英文），缺失才回退启发式。"""
    val = _en(getattr(outline, "reading_strategy", ""))
    if val:
        return val
    return "Asking Questions (KWL)" if is_nf else "Making Connections: Text to Self"


def _reading_skill(outline: BookOutline, is_nf: bool) -> str:
    val = _en(getattr(outline, "reading_skill", ""))
    if val:
        return val
    return "Main Idea and Key Details" if is_nf else "Story Elements"


def _cefr_default(level: str) -> str:
    key = str(level or "").strip().lower()
    if "smart" in key:
        return "Pre-A1"
    digits = "".join(ch for ch in key if ch.isdigit())
    mapping = {0: "Pre-A1", 1: "Pre-A1", 2: "A1", 3: "A1+", 4: "A2", 5: "B1", 6: "B1+"}
    try:
        return mapping.get(int(digits), "A1")
    except ValueError:
        return "A1"


# ============================================================
#  DOCX 工具
# ============================================================
def _set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = FONT_EN
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_EN)
    rfonts.set(qn("w:hAnsi"), FONT_EN)
    rfonts.set(qn("w:eastAsia"), FONT_CN)


def _set_a4_margins(doc: Document) -> None:
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)


def _title_block(doc, outline: BookOutline, is_nf: bool) -> None:
    """文档抬头：大蓝标题 + 书名 + 灰副标 + 蓝分隔线（对齐 SOP 抬头）。"""
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _font(cap.add_run("TEACHER'S GUIDE"), size_pt=11, bold=True, color=C_GRAY)

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _font(h.add_run(capitalize_names(outline.title)), size_pt=26, bold=True, color=C_BLUE)

    genre = "Nonfiction" if is_nf else "Fiction"
    theme = (outline.theme or "").strip() or _theme_fallback(outline, is_nf)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    book_no = (getattr(outline, "book_number", "") or "").strip() or "1"
    _font(sub.add_run(f"Level {_level_label(outline.level)}  \u2022  Book {book_no}  \u2022  {genre}  \u2022  {theme}"),
          size_pt=12, italic=True, color=C_GRAY)
    _divider(doc)


def _divider(doc, color: str = C_BLUE) -> None:
    """一条细蓝分隔线（段落下边框）。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _heading(doc, text: str, level: int = 2) -> None:
    """彩色标题：Part(2)=主蓝大号+下分隔线；小节(3)=深蓝；页/活动(4)=深蓝中号。"""
    p = doc.add_paragraph()
    try:
        p.style = f"Heading {min(max(int(level), 1), 4)}"
    except Exception:
        pass
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10 if level <= 2 else 7)
    p.paragraph_format.space_after = Pt(3)
    sizes = {1: 24, 2: 17, 3: 13.5, 4: 12}
    colors = {1: C_BLUE, 2: C_BLUE, 3: C_BLUE_DK, 4: C_BLUE_DK}
    run = p.add_run(capitalize_names(text))
    _font(run, size_pt=sizes.get(level, 12), bold=True, color=colors.get(level, C_BLUE_DK))
    if level <= 2:
        _divider(doc)


def _para(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _font(p.add_run(capitalize_names(text)), size_pt=11, color=C_INK)


def _para_italic(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _font(p.add_run(capitalize_names(text)), size_pt=10.5, italic=True, color=C_GRAY)


def _tagged(doc, tag: str, body: str) -> None:
    """加粗 tag + 正常正文。若 tag 是 [L0]/[L1]/[L2]，按层级上色。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tag_disp = capitalize_names(tag)
    tag_key = tag.strip().strip("[]").upper()
    tag_color = TAG_COLOR.get(tag_key)
    r_tag = p.add_run(tag_disp + " ")
    _font(r_tag, size_pt=11, bold=True, color=tag_color or C_INK)
    _font(p.add_run(capitalize_names(body)), size_pt=11, color=C_INK)


def _callout(doc, lines, *, bg: str = BG_BLUE, header: str | None = None,
             header_bg: str = C_BLUE) -> None:
    """彩色提示框：可选彩色头条 + 阴影正文格（单列表格实现）。

    lines: str 或 list[str|tuple]。tuple=(tag, body) → 加粗 tag。
    """
    if isinstance(lines, str):
        lines = [lines]
    nrows = (1 if header else 0) + 1
    table = doc.add_table(rows=nrows, cols=1)
    table.autofit = True
    ri = 0
    if header:
        hc = table.rows[0].cells[0]
        _shade_cell(hc, header_bg)
        hc.text = ""
        hp = hc.paragraphs[0]
        _font(hp.add_run(capitalize_names(header)), size_pt=11.5, bold=True, color="FFFFFF")
        ri = 1
    body_cell = table.rows[ri].cells[0]
    _shade_cell(body_cell, bg)
    body_cell.text = ""
    first = True
    for ln in lines:
        p = body_cell.paragraphs[0] if first else body_cell.add_paragraph()
        first = False
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if isinstance(ln, tuple):
            tag, body = ln
            _font(p.add_run(capitalize_names(tag) + " "), size_pt=10.5, bold=True, color=C_INK)
            _font(p.add_run(capitalize_names(body)), size_pt=10.5, color=C_INK)
        else:
            _font(p.add_run(capitalize_names(ln)), size_pt=10.5, color=C_INK)
    _set_table_borders(table, BORDER, sz=4)


def _kv_table(doc, rows: list[tuple[str, str]]) -> None:
    """两列信息表：蓝底白字表头(Field/Content) + 斑马纹 + 加粗键列。"""
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.columns[0].width = Cm(4.5)
    table.columns[1].width = Cm(12.5)
    _cell_text(table.rows[0].cells[0], "Field", bold=True, color="FFFFFF", fill=C_BLUE)
    _cell_text(table.rows[0].cells[1], "Content", bold=True, color="FFFFFF", fill=C_BLUE)
    for i, (k, v) in enumerate(rows):
        zebra = BG_BLUE if i % 2 == 0 else BG_WHITE
        _cell_text(table.rows[i + 1].cells[0], k, bold=True, color=C_BLUE_DK, fill=zebra)
        _cell_text(table.rows[i + 1].cells[1], v, bold=False, color=C_INK, fill=zebra)
    _set_table_borders(table, BORDER, sz=4)


def _grid_table(doc, header: list[str], body: list[list[str]], *, widths: list[float] | None = None) -> None:
    """网格表：蓝底白字表头 + 斑马纹 + 细边框；[L0/L1/L2] 标签自动上色；分节行灰底。"""
    cols = len(header)
    table = doc.add_table(rows=1 + len(body), cols=cols)
    if widths:
        for ci, w in enumerate(widths[:cols]):
            for row in table.rows:
                row.cells[ci].width = Cm(w)
    for ci, htext in enumerate(header):
        _cell_text(table.rows[0].cells[ci], htext, bold=True, color="FFFFFF", fill=C_BLUE)
    zi = 0
    for ri, rowvals in enumerate(body, start=1):
        is_section = len(rowvals) >= 2 and rowvals[1] == "" and rowvals[-1] == ""
        if is_section:
            fill = BG_GRAY
        else:
            fill = BG_BLUE if zi % 2 == 0 else BG_WHITE
            zi += 1
        for ci in range(cols):
            val = rowvals[ci] if ci < len(rowvals) else ""
            tag_key = val.strip().strip("[]").upper()
            tcol = TAG_COLOR.get(tag_key) if ci == 0 else None
            _cell_text(table.rows[ri].cells[ci], val,
                       bold=is_section or bool(tcol),
                       color=tcol or C_INK, fill=fill)
    _set_table_borders(table, BORDER, sz=4)


def _cell_text(cell, text: str, *, bold: bool, color: str = C_INK, fill: str | None = None) -> None:
    if fill:
        _shade_cell(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _font(p.add_run(capitalize_names(text)), size_pt=10.5, bold=bold, color=color)


def _shade_cell(cell, fill_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _set_table_borders(table, color: str, sz: int = 4) -> None:
    """给整张表设统一细边框（淡色）。"""
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def _font(run, *, size_pt: float, bold: bool = False, italic: bool = False,
          color: str | None = None) -> None:
    run.font.name = FONT_EN
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_EN)
    rfonts.set(qn("w:hAnsi"), FONT_EN)
    rfonts.set(qn("w:eastAsia"), FONT_CN)


def _level_label(level: str) -> str:
    s = (level or "").strip()
    if not s:
        return "1"
    if s.lower().startswith("smart"):
        return "Smart"
    digits = "".join(ch for ch in s if ch.isdigit())
    return digits or s
