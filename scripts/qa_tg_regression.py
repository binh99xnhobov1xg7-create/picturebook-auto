"""Regression checks for L3-L4 Teacher's Guide generation.

Run from repo root:
    python scripts/qa_tg_regression.py
"""
from __future__ import annotations

from pathlib import Path
import sys
import tempfile

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parent))

from parser import BookOutline, PageSpec  # noqa: E402
from teacher_guide_builder import build_teacher_guide  # noqa: E402
from tg_quality import TGQualityError  # noqa: E402
from tg_page_overrides import official_page_text_override  # noqa: E402


def _outline(title: str, level: str, book: str, texts: list[str], **kwargs) -> BookOutline:
    pages = [PageSpec(index=0, page_type="cover", text="")]
    for i, text in enumerate(texts, 1):
        pages.append(PageSpec(index=i, page_type="story", text=text))
    return BookOutline(title=title, pages=pages, level=level, book_number=book, **kwargs)


def _mia_outline(texts: list[str] | None = None) -> BookOutline:
    texts = texts or [
        "Mia has many things to do this week.",
        "She has homework to finish. Her room is messy.",
        "The piano show is on Sunday. Mia wants to play well.",
        "Mia feels worried. She makes a seven-day plan.",
        "She will do homework first. She will clean her room on Tuesday.",
        "She will practice the piano every day. She will play for one hour a day.",
        "Mia is happy and proud of her plan.",
    ]
    out = _outline(
        "Mia and the Seven-Day Plan",
        "3",
        "1",
        texts,
        cefr="A1",
        lexile="210L-400L",
        vocabulary_simple=["week", "homework", "plan", "practice"],
        phonics='Long vowel "ay" as in "day"',
        grammar_focus="[Subject] will [action].",
        reader_type="Fiction",
        fiction_type="fiction",
        reading_strategy="Predicting",
        reading_skill="Sequencing Events",
        graphic_organizer="Plan Chart",
    )
    out._worksheet_second_reading_mode = "planchart"
    return out


def _scotland_outline() -> BookOutline:
    texts = [
        "Mia and Tommy visited Scotland with their family.",
        "They saw a tall stone castle by a blue loch.",
        "Men were playing the bagpipes on a green hill.",
        "They were amazed to see a man dancing in a kilt.",
        "A cheeky sheep grabbed their paper map and ran away.",
        "A nice woman found the map and gave it back.",
        "The friendly Scots made the journey unforgettable.",
    ]
    out = _outline(
        "Visiting Scotland",
        "4",
        "13",
        texts,
        cefr="A2",
        lexile="410L-600L",
        vocabulary_simple=["culture", "castle", "bagpipes", "journey"],
        phonics='Long "ee" sound as in sheep',
        grammar_focus="There was / There were + noun + prepositional phrase",
        reader_type="Fiction",
        fiction_type="fiction",
        reading_strategy="Making Connections: Text to Self",
        reading_skill="Story Elements",
    )
    return out


def _travel_outline() -> BookOutline:
    texts = [
        "Mia and Tommy are looking at travel photos from last year.",
        "There is a blue ocean with big waves and a crab. There is a tall, snowy mountain in one photo.",
        "There is a wide, hot desert with soft sand.",
        "There is a deep, rocky canyon with tall walls.",
        "There is a green valley full of bright flowers.",
        "One photo is missing. It is the one they took at the Grand Canyon. Where could it be?",
        'They ask Mom, and she says, "It was mailed to Grandma last week." "Oh, yes!" They both laugh.',
    ]
    return _outline(
        "Mia and Tommy Travel the World",
        "4",
        "1",
        texts,
        cefr="A2",
        lexile="410L-600L",
        vocabulary_simple=["valley", "desert", "ocean", "canyon"],
        phonics='Long "o" sound as in photo',
        grammar_focus="There is / There are + adjective + noun",
        reader_type="Fiction",
        fiction_type="fiction",
        reading_strategy="Predicting",
        reading_skill="Story Elements",
    )


def _doc_text(path: Path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_l3_mia_pass(tmp: Path) -> None:
    path = tmp / "mia_tg.docx"
    build_teacher_guide(_mia_outline(), path)
    text = _doc_text(path)
    required = [
        "Level 3",
        "Book 1",
        "Subject + will + base verb",
        "Pages 2-8",
        "first, clean her room, every day, play for one hour",
        "Problem - Plan - Result",
        "What prediction did you make about Mia's plan?",
    ]
    forbidden = [
        "Quality Checklist",
        "Mia follows her plan",
        "Mia completes all her tasks",
        "excited at the start",
        "week, homework, plan, practice\nCompletion Guidance",
    ]
    for item in required:
        assert item in text, item
    for item in forbidden:
        assert item not in text, item


def test_l3_mia_bad_pages_block(tmp: Path) -> None:
    bad = [
        "Mia has many things to do this week. She has homework to finish.",
        "Her room is messy. The piano show is on Sunday.",
        "Mia wants to play well. Mia feels worried.",
        "She makes a seven-day plan.",
        "She will do homework first. She will clean her room on Tuesday.",
        "She will practice the piano every day. She will play for one hour a day.",
        "Mia is happy and proud of her plan.",
    ]
    try:
        build_teacher_guide(_mia_outline(bad), tmp / "bad_mia_tg.docx")
    except TGQualityError as exc:
        assert "l3_mia_page_text_mismatch" in str(exc)
        return
    raise AssertionError("Bad L3-1 page split should be blocked.")


def test_l3_mia_official_story_page_override() -> None:
    outline = _mia_outline([""] * 7)
    official_story = " ".join([
        "Mia has many things to do this week.",
        "She has homework to finish.",
        "Her room is messy.",
        "The piano show is on Sunday.",
        "Mia wants to play well.",
        "Mia feels worried.",
        "She makes a seven-day plan.",
        "She will do homework first.",
        "She will clean her room on Tuesday.",
        "She will practice the piano every day.",
        "She will play for one hour a day.",
        "Mia is happy and proud of her plan.",
    ])
    chunks = official_page_text_override(outline)
    assert chunks is not None
    for i, text in enumerate(chunks, start=1):
        outline.pages[i].text = text
    assert outline.pages[1].text == "Mia has many things to do this week."
    assert outline.pages[2].text == "She has homework to finish. Her room is messy."
    assert outline.pages[3].text == "The piano show is on Sunday. Mia wants to play well."
    assert outline.pages[4].text == "Mia feels worried. She makes a seven-day plan."


def test_l4_scotland_smoke(tmp: Path) -> None:
    path = tmp / "scotland_tg.docx"
    build_teacher_guide(_scotland_outline(), path)
    text = _doc_text(path)
    assert "Teacher's Guide" in text or "TEACHER'S GUIDE" in text
    assert "Quality Checklist" not in text
    assert "L2 questions appear at the final Pause Point" not in text
    assert "Lesson Overview" in text


def test_l4_travel_dialogue_quotes_pass(tmp: Path) -> None:
    path = tmp / "travel_tg.docx"
    build_teacher_guide(_travel_outline(), path)
    text = _doc_text(path)
    assert "Page 8:" in text
    assert "It was mailed to Grandma last week." in text
    assert "Oh, yes!" in text


def main() -> None:
    with tempfile.TemporaryDirectory() as d:
        tmp = Path(d)
        test_l3_mia_pass(tmp)
        print("PASS L3-1 correct-pages TG generation")
        test_l3_mia_bad_pages_block(tmp)
        print("PASS L3-1 bad-pages blocker")
        test_l3_mia_official_story_page_override()
        print("PASS L3-1 official-story page override")
        test_l4_scotland_smoke(tmp)
        print("PASS L4B13 smoke regression")
        test_l4_travel_dialogue_quotes_pass(tmp)
        print("PASS L4-1 dialogue quote regression")


if __name__ == "__main__":
    main()
