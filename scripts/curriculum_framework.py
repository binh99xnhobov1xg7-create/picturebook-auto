"""宏观 0-6 课程大纲 → 可分享 Word 文档。

复用 teacher_guide_builder 的官方 SOP 色块/表格渲染工具，保证与现有
TG / Worksheet 交付物生态一致。数据源为 level_profiles.LEVEL_PROFILES。

生成：一个总览对比表 + 每级一节详情（CEFR/Lexile/字数/IP/知识体系/句法/
词汇/阅读技能/认知层级/题目配比/教学流程差异/主题/课堂教学法/学完产出）。
"""
from __future__ import annotations

from pathlib import Path

from docx import Document

import teacher_guide_builder as tg
from level_profiles import LEVEL_PROFILES, ORDERED_LEVELS, LevelProfile


def _mix_str(p: LevelProfile) -> str:
    a, b, c = p.question_mix
    return f"[L0] {a:.0%} · [L1] {b:.0%} · [L2] {c:.0%}"


def build_curriculum_framework(out_path: Path) -> Path:
    doc = Document()
    tg._set_default_font(doc)
    tg._set_a4_margins(doc)

    # ---- 抬头 ----
    cap = doc.add_paragraph()
    tg._font(cap.add_run("CURRICULUM FRAMEWORK"), size_pt=11, bold=True, color=tg.C_GRAY)
    h = doc.add_paragraph()
    tg._font(h.add_run("Levels 0-6 \u00b7 \u8bfe\u7a0b\u96be\u5ea6\u9636\u68af\u4e0e\u5b66\u4e60\u5730\u56fe"), size_pt=24, bold=True, color=tg.C_BLUE)
    sub = doc.add_paragraph()
    tg._font(sub.add_run("Pre-A1 \u2192 B1 \u00b7 \u4ece\u770b\u56fe\u547d\u540d\u5230\u6279\u5224\u6027\u9605\u8bfb \u00b7 \u5b98\u65b9 S&S \u5927\u7eb2\u63a8\u5bfc"),
            size_pt=12, italic=True, color=tg.C_GRAY)
    tg._divider(doc)

    # ---- 难度阶梯说明 ----
    tg._callout(
        doc,
        [
            "\u8fd9\u6761\u9636\u68af\u662f\u6574\u5957\u4f53\u7cfb\u7684\u5e95\u5c42\u903b\u8f91\uff1a\u8bcd\u6c47\u6df1\u5ea6\u3001\u53e5\u6cd5\u3001\u89e3\u7801/\u6784\u8bcd\u3001\u9605\u8bfb\u6280\u80fd\u3001\u63d0\u95ee\u8ba4\u77e5\u5c42\u7ea7\u540c\u6b65\u4e0a\u5347\u3002",
            ("\u5206\u6c34\u5cad\u4e00 (L2\u2192L3):", "\u9884\u6559\u8bcd + 4\u6b65 Phonics + \u5168\u9875 Picture Walk \u2192 \u8bed\u5883\u5185\u5b66\u8bcd + 2\u6b65 Phonics \u610f\u8bc6 + \u4f53\u88c1\u5b9a\u8303\u56f4 + \u52a8\u6001 Pause Point\u3002"),
            ("\u5206\u6c34\u5cad\u4e8c (L4\u2192L5):", "\u81ea\u7136\u62fc\u8bfb (Phonics) \u2192 \u6784\u8bcd\u6cd5 (Word Formation)\uff1b\u6838\u5fc3\u8bcd 4\u21925\uff1b\u6587\u672c ~90 \u8bcd \u2192 ~130-200 \u8bcd\u3002"),
        ],
        bg=tg.BG_BLUE, header="\u2605 0-6 \u96be\u5ea6\u9636\u68af\u603b\u8bf4", header_bg=tg.C_BLUE,
    )

    # ---- 总览对比表 ----
    tg._heading(doc, "\u603b\u89c8\u5bf9\u6bd4")
    header = ["Level", "CEFR", "Lexile", "\u5b57\u6570", "\u77e5\u8bc6\u4f53\u7cfb", "\u9605\u8bfb\u6280\u80fd / \u8ba4\u77e5"]
    body = []
    for lvl in ORDERED_LEVELS:
        p = LEVEL_PROFILES[lvl]
        body.append([
            f"L{lvl}", p.cefr, p.lexile, p.word_count,
            p.decoding_system,
            f"{p.reading_skill_tier}\uff08{p.cognitive_demand}\uff09",
        ])
    tg._grid_table(doc, header, body, widths=[1.4, 2.0, 2.6, 2.2, 4.2, 4.6])

    # ---- 每级详情 ----
    for lvl in ORDERED_LEVELS:
        p = LEVEL_PROFILES[lvl]
        tg._heading(doc, f"Level {lvl}  \u00b7  {p.cefr}  \u00b7  {_band_cn(p.band)}")

        tg._kv_table(doc, [
            ("CEFR / Lexile", f"{p.cefr}  \u00b7  {p.lexile}"),
            ("\u5355\u672c\u5b57\u6570", p.word_count),
            ("IP \u5f62\u8c61\u5e74\u9f84", f"{p.ip_age} \u5c81"),
            ("\u77e5\u8bc6\u4f53\u7cfb", f"{p.decoding_system} \u2014 {p.decoding_detail}"),
            ("\u53e5\u6cd5\u91cd\u70b9", p.syntax_focus),
            ("\u6838\u5fc3\u8bcd\u6c47", f"\u6bcf\u672c {p.vocab_count} \u8bcd\uff08{p.vocab_model}\uff09\uff1b{p.vocab_teaching}"),
            ("\u9605\u8bfb\u6280\u80fd", p.reading_skill_tier),
            ("\u8ba4\u77e5\u5c42\u7ea7", p.cognitive_demand),
            ("\u9898\u76ee\u96be\u5ea6\u914d\u6bd4", _mix_str(p)),
            ("\u6559\u5b66\u6d41\u7a0b", f"Picture Walk: {p.picture_walk} \u00b7 Pause: {p.pause_points} \u00b7 "
                                     f"Phonics {p.phonics_steps} \u6b65" + ("\u00b7 \u8001\u5e08\u6307\u8bfb" if p.finger_tracking else "")),
        ])

        tg._heading(doc, "\u8bcd\u6c47\u4e3b\u9898", level=3)
        tg._para(doc, "\u3001".join(p.themes))

        tg._heading(doc, "\u8bfe\u5802\u8bb2\u89e3\u65b9\u5f0f", level=3)
        for mv in p.teaching_moves:
            tg._para(doc, f"\u2022 {mv}")

        tg._heading(doc, "\u5b66\u5b8c\u80fd\u5f97\u5230\u4ec0\u4e48", level=3)
        tg._tagged(doc, "\u5f97\u5230\uff08\u6210\u679c\uff09:", p.outcome_get)
        tg._tagged(doc, "\u5b66\u5230\uff08\u77e5\u8bc6/\u6280\u80fd\uff09:", p.outcome_learn)
        tg._tagged(doc, "\u7ec3\u5230\uff08\u8fc1\u79fb\u5e94\u7528\uff09:", p.outcome_practice)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def _band_cn(band: str) -> str:
    return {"low": "\u542f\u8499\u6863 (L0-2)", "mid": "\u63d0\u5347\u6863 (L3-4)", "high": "\u8fdb\u9636\u6863 (L5-6)"}.get(band, band)


if __name__ == "__main__":
    out = Path(__file__).resolve().parents[1] / "outputs" / "_framework" / "Curriculum_Framework_L0-L6.docx"
    print("WROTE", build_curriculum_framework(out))
