"""Reading Report DOCX 生成器 v2.1 — 严格 1 页 + 顶格 + 课堂参与度方框。

排版完全复刻 VIPKID 官方模板（Desktop/L0_Book92/L0_Book92 reading report.docx
+ Desktop/LX_BookXX reading report Sample.docx），并满足以下口径：

  • 严格 1 页 A4 portrait（A4 portrait 8.27×11.69 in，margins 1.0 cm）
  • 表格 cell 段落顶格对齐（vertical = TOP）
  • 字号 11pt + 行距 1.2（紧凑但不挤）
  • 阅读表达星级用 ★ (U+2605) + 橙色 #E97A24（emoji ⭐ 在 Word/WPS 上会渲染成 *）
  • 课堂参与度 emoji 紧贴 label，每组后跟一个空白方框 ☐ 供学生打勾

题量梯度（与 picture-book-workflow QA 规则一致）:
  L0/L1/L2/Smart = 4 题, ★ + ★★ + ★★ + ★★★
  L3/L4/L5/L6    = 5 题, ★ + ★★ + ★★ + ★★ + ★★★
末尾 ★★★ 题为生活化拓展题（不带 (P#)），其余必须带 (P#)。
"""
from __future__ import annotations

import math
import re
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

import config
from parser import BookOutline
from text_format import _to_us_spelling, capitalize_names, format_word_answer, format_sentence_answer


FONT_EN = "Poppins"
FONT_CN = "阿里巴巴普惠体 2.0 55 Regular"
EMOJI_FONT = "Segoe UI Emoji"
# 符号字体：★ □ 等几何符号统一用 Arial 渲染（Poppins/CJK 缺字会变 tofu/乱码）
SYMBOL_FONT = "Arial"
EMPTY_BOX = "\u25A1"  # □ WHITE SQUARE（比 ☐ U+2610 兼容性更好）


def _strip_wrapping_quotes(text: str) -> str:
    """去掉整段被引号包裹的情况（QA：主体文本粘贴时不要带引号）。

    仅当整段以引号开头且以引号结尾时去壳；保留句内正常对话引号。
    """
    s = (text or "").strip()
    pairs = {'"': '"', "'": "'", "\u201c": "\u201d", "\u2018": "\u2019",
             "\u300c": "\u300d", "\u00ab": "\u00bb"}
    changed = True
    while changed and len(s) >= 2:
        changed = False
        for lq, rq in pairs.items():
            if s.startswith(lq) and s.endswith(rq) and len(s) >= 2:
                s = s[1:-1].strip()
                changed = True
                break
    return s


def _normalize_student_punctuation(text: str) -> str:
    """Use standard ASCII quotes/apostrophes/hyphens in student-facing text."""
    out = (
        str(text or "")
        .replace("\u201c", '"').replace("\u201d", '"')
        .replace("\u2018", "'").replace("\u2019", "'")
        .replace("\u02bc", "'")
        .replace("\uff02", '"').replace("\uff07", "'")
        .replace("\ufffe", "-").replace("\u00ad", "-")
        .replace("\u2010", "-").replace("\u2011", "-")
        .replace("\u2012", "-").replace("\u2013", "-").replace("\u2014", "-")
    )
    out = re.sub(r"(?<=[A-Za-z])\?(?=s\b)", "'", out)
    out = re.sub(r"\?([A-Za-z]{1,12})\?", r'"\1"', out)
    return out

# 表格列宽（严格按 sample，dxa 即 1/20 pt）
COL_WIDTHS_DXA = [1428, 1307, 1293, 1246, 1369, 1308, 1292, 1243]
TABLE_WIDTH_DXA = sum(COL_WIDTHS_DXA)  # 10486

# 行高（twips，1440 twips = 1 inch）。设为 atLeast，由内容撑开但有最小值
# 经实测：A4 portrait 可用高约 10.98"，下列值能让 L0-L6 都恰好落在一页内
HEADER_ROW_TWIPS = 480        # 灰底 section header
VOCAB_ROW_TWIPS = 700
DIFF_ROW_TWIPS = 1500         # 4 行 11pt 标签/值（标题/字数/CEFR/语法）
PHONICS_ROW_TWIPS = 600
FLUENCY_ROW_TWIPS = 1300      # 故事正文行（atLeast，长内容自然撑开）
QUESTIONS_ROW_TWIPS = 2200    # 4-5 题（atLeast，长题自然撑开）
ENGAGE_ROW_TWIPS = 900

HEADER_FILL = "F1F1F1"

# 字号 / 行距口径（pt）—— 对齐官方 RR Sample 模板（标题17 / 章节头14左对齐 / 难度14）
TITLE_PT = 17
NAME_PT = 11
SECTION_PT = 14          # 章节标题（官方模板：Poppins 14pt 粗体，左对齐）
DIFF_PT = 14             # 阅读难度内容（类型/字数/词汇难度/语法难度，官方 14pt）
BODY_PT = 11             # 正文/题目/词汇（紧凑，保证 1 页）
LINE_SPACING = 1.2
PARA_AFTER_PT = 1   # 段后空隙（题目 / 难度 / 故事正文统一）

# 星级颜色（橙红，对照 sample）
STAR_COLOR = "E97A24"

# Logo 资源（可被 config.BRAND_DIR 下的同名文件覆盖）
LOGO_PATH = config.BRAND_DIR / "dino_reading_logo.png"


# ============================================================================
# 自适应排版规划器 —— 严格 1 页的核心
# ----------------------------------------------------------------------------
# 思路：A4 portrait 可用高度固定，按"实际内容"估算每个 section 撑开的高度，
# 从大字号档位逐档下探，选第一个能整体落进一页预算的档位（字号/行距/行高/段后）。
# 这样无论故事多长、题目多长、语法行多长，都能自动收进一页而不溢出。
# ============================================================================
def _plan_layout(outline, shrink_steps: int = 0) -> dict:
    PAGE_H = 16838           # 29.7cm A4 高 (twips)
    MARGINS = 1020           # 上下各 0.9cm
    TITLE_BLOCK = 1520       # 标题行（含 logo）估高；按长书名折成两行预留（防溢出第二页）
    NAME_LINE = 380          # 姓名/日期行
    SAFETY = 1260            # 安全余量加大：RR 必须严格压在 1 页内
    budget = PAGE_H - MARGINS - TITLE_BLOCK - NAME_LINE - SAFETY

    def cpl(pt: float) -> int:
        """满表宽下每行英文字符数（保守略小 → 多估行数，避免溢出）。"""
        return max(20, int(960 / pt))

    def line_h(pt: float, ls: float) -> int:
        return int(pt * 20 * ls * 1.16)

    grammar = _clean_rr_text(_grammar_difficulty_text(outline) or "—")
    passage = _strip_wrapping_quotes(" ".join(
        _strip_wrapping_quotes((p.text or "").strip())
        for p in outline.pages
        if p.page_type == "story" and p.text and p.text.strip()
    ))
    _cap = getattr(outline, "_rr_passage_char_cap", None)
    if _cap and len(passage) > _cap:
        passage = passage[:_cap]
    title = outline.title or ""
    p_chars = len(passage) + len(title) + 2

    qs = _resolve_rr_questions(outline)
    q_lens: list[int] = []
    for q in qs:
        # 序号"1. " + 句末"?" ≈ 4；(P#) ≈ 6；★ 星级 ≈ 4（保证一行不被星星挤换行）
        extra = 4 + 4 + (6 if q.get("page") is not None else 0)
        q_lens.append(len((q.get("q") or "")) + extra)

    phon = (_normalize_morphology(outline.phonics, outline)
            if _is_morphology_level(outline.level)
            else _normalize_phonics(outline.phonics)) or ""
    vocab_ok = any(_vocab_words_for_rr(outline))

    # 对齐人工模板（L5-1 5.27）：正文/难度/章节标题统一 12pt（标题 14pt 在别处），
    # 行距适中、学生看着舒服。再往下是兜底小档，仅在内容超多时才用。
    profiles = [
        dict(body=11.5, diff=11.5, sec=11.5, ls=1.06, pa=0),
        dict(body=11.0, diff=11.0, sec=11.0, ls=1.04, pa=0),
        dict(body=10.5, diff=10.5, sec=10.8, ls=1.02, pa=0),
        dict(body=10.0, diff=10.5, sec=10.5, ls=1.00, pa=0),
        dict(body=9.8,  diff=10.0, sec=10.2, ls=0.98, pa=0),
        dict(body=9.5,  diff=10.0, sec=10.0, ls=0.96, pa=0),
        dict(body=9.2,  diff=9.8,  sec=9.8,  ls=0.94, pa=0),
        # 兜底：更小档，确保内容再多也能压进一页
        dict(body=9.0,  diff=9.5,  sec=9.5,  ls=0.94, pa=0),
        dict(body=8.5,  diff=9.0,  sec=9.2,  ls=0.92, pa=0),
        dict(body=8.0,  diff=8.8,  sec=9.0,  ls=0.90, pa=0),
    ]

    longest_q = max(q_lens) if q_lens else 0

    def estimate(prof: dict):
        body, diff, sec, ls, pa = prof["body"], prof["diff"], prof["sec"], prof["ls"], prof["pa"]
        pa_tw = int(pa * 20)
        lh_b, lh_d = line_h(body, ls), line_h(diff, ls)

        header_h = max(420, line_h(sec, 1.0) + 150)
        headers = 6 * header_h

        g_eff = len("语法难度：") * 2 + len(grammar)   # 中文标签约占 2x 宽
        g_lines = max(1, math.ceil(g_eff / cpl(diff)))
        diff_h = max((3 + g_lines) * lh_d + 4 * pa_tw, lh_d + 80)

        # 词汇掌握 = 紧凑单行（词格 + 书写空格交替），对齐官方模板 R3≈539 twips：
        # 就一行，不拉高留白（用户拍板 2026-06-04 + 5 份官方样板）。
        vocab_h = 560 if vocab_ok else 420

        ph_lines = max(1, math.ceil((len(phon) + 10) / cpl(body)))
        phon_h = max(ph_lines * lh_b + 100, 360)

        fl_lines = 1 + max(1, math.ceil(p_chars / cpl(body)))
        fluency_h = fl_lines * lh_b + 80

        q_lines = sum(max(1, math.ceil(L / cpl(body))) for L in q_lens) or 1
        questions_h = q_lines * lh_b + len(q_lens) * pa_tw + 80

        engage_h = max(lh_b + 120, 580)

        total = headers + diff_h + vocab_h + phon_h + fluency_h + questions_h + engage_h
        mins = dict(header_h=int(header_h), diff_h=int(diff_h), vocab_h=int(vocab_h),
                    phon_h=int(phon_h), fluency_h=int(fluency_h),
                    questions_h=int(questions_h), engage_h=int(engage_h))
        return total, mins

    # 先收集所有"能塞进一页"的档；在其中优先选"每道阅读题都能一行"的最大字号档（C4）。
    fitting: list[tuple[dict, dict]] = []
    for prof in profiles:
        total, mins_ = estimate(prof)
        if total <= budget:
            fitting.append((prof, mins_))

    # 选档原则（对齐官方模板）：优先用最大可一页容纳的字号（≈12pt，和官方一致），
    # 允许长题目自然换行——官方样板里长题本就折行，不为"每题一行"硬压小字号。
    chosen, mins = None, None
    if fitting:
        chosen, mins = fitting[0]            # 最大可容纳字号
        # 若最大字号下长题需折行，而仅小半档即可让每题一行且字号仍 >=11pt，则取该档（更美观）
        if longest_q > cpl(chosen["body"]):
            for prof, mins_ in fitting:
                if prof["body"] >= 11.0 and longest_q <= cpl(prof["body"]):
                    chosen, mins = prof, mins_
                    break
    if chosen is None:                       # 极端兜底：用最小档（仍强制一页优先）
        chosen = profiles[-1]
        _, mins = estimate(profiles[-1])

    # —— 降档重排（硬保证 1 页）——
    # build_reading_report 实测渲染页数 >1 时，会带着递增的 shrink_steps 重排：
    # 在已选档基础上强制下移 N 档（更小字号/更紧行距），直到实测落进一页。
    if shrink_steps > 0:
        try:
            cur_idx = profiles.index(chosen)
        except ValueError:
            cur_idx = 0
        new_idx = min(len(profiles) - 1, cur_idx + shrink_steps)
        chosen = profiles[new_idx]
        _, mins = estimate(chosen)

    # —— 整页铺满（用户拍板 2026-06-04：内容必须铺满竖版 A4）——
    # 把剩余高度几乎全部补给两块"书写/阅读区"：答题区(阅读表达，主区，占大头) + 流利度(正文区)，
    # 二者顶部对齐 → 内容在上、下方留白 = 学生书写空间；只保留极小安全余量，整页填满不空。
    used = (6 * mins["header_h"] + mins["diff_h"] + mins["vocab_h"] + mins["phon_h"]
            + mins["fluency_h"] + mins["questions_h"] + mins["engage_h"])
    leftover = budget - used
    # 降档 4 档及以上 = 正在和溢出搏斗，彻底关闭"铺满"扩张（绝不再加高度）。
    if leftover > 200 and shrink_steps < 2:
        # 余量按三大内容区（难度/正文/答题）当前高度比例均摊 → 整页均衡铺满、
        # 不在某一块堆出大片空白（对齐官方模板均衡的版面），词汇/拼读/参与度保持紧凑。
        fill_keys = ["diff_h", "fluency_h", "questions_h"]
        base = sum(mins[k] for k in fill_keys) or 1
        # 铺满系数（用户拍板 2026-06-06：底部不能留大片空白）：首轮吃掉 ~92% 余量，
        # 降档重排时每档少给 12%（防被撑大又溢出到第二页）。
        fill_factor = max(0.15, 0.35 - 0.12 * shrink_steps)
        give = int(leftover * fill_factor)   # 留安全余量，严格防溢出到第二页
        natural = {k: max(1, mins[k]) for k in fill_keys}
        for k in fill_keys:
            mins[k] += int(give * mins[k] / base)
        # 关键（用户拍板 2026-06-06）：竖版 A4 内容少时，靠【放大文字行距】让文字本身铺满
        # 加高后的格子，而不是把文字顶在上面、下方留白。按 加高后/自然 比例放大行距，分段封顶。
        # （封顶偏保守 + count_pages 实测降档双保险，严格守 1 页。）
        base_ls = float(chosen["ls"])
        for _key, _ls_name, _cap in (
            ("diff_h", "ls_diff", 1.8),
            ("fluency_h", "ls_fluency", 1.9),
            ("questions_h", "ls_questions", 2.4),
        ):
            _boost = mins[_key] / natural[_key]
            chosen[_ls_name] = round(min(_cap, max(base_ls, base_ls * _boost)), 2)
    plan = dict(chosen)
    plan.update(mins)
    plan["grammar"] = grammar
    return plan


def _lp(outline) -> dict:
    return getattr(outline, "_rr_lp", {}) or {}


# ============================================================================
# 公开入口
# ============================================================================
def _render_rr_doc(outline: BookOutline, shrink_steps: int = 0):
    """按给定 shrink_steps 规划版式并构建 RR 文档（不落盘）。"""
    setattr(outline, "_rr_lp", _plan_layout(outline, shrink_steps=shrink_steps))
    doc = Document()
    _set_a4_portrait(doc)
    _set_default_style(doc)

    # 移除 Document() 自带的初始空段
    body = doc.element.body
    for p in list(body.iter(qn("w:p"))):
        body.remove(p)

    _build_title_paragraph(doc, outline)
    _build_name_date_paragraph(doc)
    _build_main_table(doc, outline)
    return doc


# 降档重排最大尝试次数（profiles 共 9 档，足够把任何内容压进一页）
_MAX_SHRINK_STEPS = 8


def build_reading_report(outline: BookOutline, out_path: Path,
                         *, with_answers: bool = False) -> Path:
    """生成阅读报告，**硬保证严格 1 张 A4**（用户拍板 2026-06-05）。

    with_answers=True → 示例答案版（演示）；False → 空白版（教师手填）。

    单页保证机制：先按自适应规划器排版并落盘，随后用 LibreOffice 实测渲染页数；
    若 >1 页，则带着递增的 shrink_steps（强制更小字号/更紧行距）重排重存，循环直到
    实测正好 1 页。soffice 不可用时回退到纯估算（规划器本身已偏保守）。
    """
    setattr(outline, "_rr_with_answers", bool(with_answers))
    out_path.parent.mkdir(parents=True, exist_ok=True)

    # 先按最佳（最大可容纳）字号渲染落盘
    doc = _render_rr_doc(outline, shrink_steps=0)
    doc.save(str(out_path))

    # 实测页数 → 超一页则降档重排，硬保证 1 页
    try:
        from doc_preview import count_pages
    except Exception:
        count_pages = None  # type: ignore

    if count_pages is not None:
        pages = count_pages(out_path)
        shrink = 0
        # pages 为 None = 无法实测（soffice/fitz 缺失）→ 不强制，回退估算
        while pages is not None and pages > 1 and shrink < _MAX_SHRINK_STEPS:
            shrink += 1
            doc = _render_rr_doc(outline, shrink_steps=shrink)
            doc.save(str(out_path))
            pages = count_pages(out_path)

        # 最终硬兜底：极端超长内容连最小档仍溢出 → 逐步截断流利度区原文，
        # 直到实测正好 1 页（杜绝任何情况下出现第二张 A4）。
        if pages is not None and pages > 1:
            full_len = len(" ".join(
                (p.text or "").strip() for p in outline.pages
                if p.page_type == "story" and (p.text or "").strip()
            ))
            cap = max(280, int(full_len * 0.85))
            while pages is not None and pages > 1 and cap >= 200:
                setattr(outline, "_rr_passage_char_cap", cap)
                doc = _render_rr_doc(outline, shrink_steps=_MAX_SHRINK_STEPS)
                doc.save(str(out_path))
                pages = count_pages(out_path)
                cap = int(cap * 0.85)
            setattr(outline, "_rr_passage_char_cap", None)

    return out_path


# ============================================================================
# 标题 / 姓名段
# ============================================================================
def _build_title_paragraph(doc, outline: BookOutline) -> None:
    """标题行：左侧书名（17pt 粗体），右侧 VIPKID Dino logo —— 用无边框两列表布局，
    保证长书名不会把 logo 挤到下一行（对齐官方模板：标题左、logo 右上）。"""
    title_str = f"阅读报告 {_level_label(outline.level)} - {capitalize_names(outline.title or '')}"

    # 标题尽量一行：按有效字宽（CJK 记 2）反推字号，长标题自动缩小（用户拍板）。
    # L3-6（高级别）页首标题统一四号 14pt 上限（用户反馈 2026-06-08）；其余沿用 TITLE_PT 上限。
    left_ratio = 0.88
    left_w_in = (TABLE_WIDTH_DXA * left_ratio) / 1440.0
    eff_len = sum(2 if ord(ch) > 0x2E80 else 1 for ch in title_str)
    high = _is_high_level(outline.level)
    candidates = (13, 12.5, 12, 11.5, 11) if high else (TITLE_PT, 16, 15, 14, 13, 12)
    title_pt = candidates[0]
    for cand in candidates:
        cap = left_w_in / (cand * 0.0085)
        if eff_len <= cap:
            title_pt = cand
            break
    else:
        title_pt = 13

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 表格整体页面居中（用户反馈 2026-06-08）
    _set_tbl_w(table, TABLE_WIDTH_DXA)
    _set_tbl_grid(table, [int(TABLE_WIDTH_DXA * left_ratio), TABLE_WIDTH_DXA - int(TABLE_WIDTH_DXA * left_ratio)])
    _set_tbl_no_borders(table)

    left = table.rows[0].cells[0]
    right = table.rows[0].cells[1]
    _vert_center(left)
    _vert_center(right)

    _clear_paragraphs(left)
    pl = left.paragraphs[0]
    pl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pl.paragraph_format.space_before = Pt(0)
    pl.paragraph_format.space_after = Pt(0)
    pl.paragraph_format.line_spacing = 1.0
    run = pl.add_run(title_str)
    _bind_run(run, ascii_font=FONT_EN, east_asia=FONT_CN, size_pt=title_pt, bold=True)

    if LOGO_PATH.exists():
        _clear_paragraphs(right)
        pr = right.paragraphs[0]
        pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pr.paragraph_format.space_before = Pt(0)
        pr.paragraph_format.space_after = Pt(0)
        logo_run = pr.add_run()
        logo_run.add_picture(str(LOGO_PATH), width=Cm(4.6))


def _build_name_date_paragraph(doc) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.0
    pf.right_indent = Cm(0.4)

    run = p.add_run("姓名: __________  日期: _________ 年 _________ 月 _________ 日")
    _bind_run(run, ascii_font="Arial Black", east_asia=FONT_CN, size_pt=NAME_PT, bold=True)


# ============================================================================
# 主表（12 行 × 8 列）
# ============================================================================
def _vocab_grid_widths(n_words: int) -> list[int]:
    """词汇行的列宽：每词 = [词格(宽)] + [书写空格(窄)] 交替，总宽 = 表宽。
    对齐官方模板（L5-1）：词格 ≈ 1.45× 空格。"""
    n = max(1, n_words)
    cols = n * 2
    unit = TABLE_WIDTH_DXA // n
    word_w = int(unit * 0.59)
    blank_w = unit - word_w
    widths: list[int] = []
    for _ in range(n):
        widths += [word_w, blank_w]
    widths[-1] = TABLE_WIDTH_DXA - sum(widths[:-1])  # 末列吸收取整误差
    return widths


def _build_main_table(doc, outline: BookOutline) -> None:
    # 列数 = 词汇数 × 2（词格/空格交替），其余行整体合并 → 完全对齐官方模板单表结构。
    _words = [w for w in _vocab_words_for_rr(outline) if w] or [""]
    n_words = len(_words)
    cols = n_words * 2
    grid_widths = _vocab_grid_widths(n_words)

    table = doc.add_table(rows=12, cols=cols)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 主表整体页面居中（用户反馈 2026-06-08）

    _set_tbl_w(table, TABLE_WIDTH_DXA)
    _set_tbl_borders(table)
    _set_tbl_grid(table, grid_widths)

    lp = _lp(outline)
    sec_pt = lp.get("sec", SECTION_PT)
    header_h = lp.get("header_h", HEADER_ROW_TWIPS)

    # v1.8.3：L5-L6 把"自然拼读"行替换为"构词法"行
    phonics_label = "构词法" if _is_morphology_level(outline.level) else "自然拼读"
    sections = [
        ("阅读难度", lp.get("diff_h", DIFF_ROW_TWIPS), _fill_difficulty),
        ("词汇掌握", lp.get("vocab_h", VOCAB_ROW_TWIPS), _fill_vocab),
        (phonics_label, lp.get("phon_h", PHONICS_ROW_TWIPS), _fill_phonics),
        ("阅读流利度", lp.get("fluency_h", FLUENCY_ROW_TWIPS), _fill_fluency),
        ("阅读表达", lp.get("questions_h", QUESTIONS_ROW_TWIPS), _fill_questions),
        ("课堂参与度", lp.get("engage_h", ENGAGE_ROW_TWIPS), _fill_engagement),
    ]

    for sec_idx, (label, content_h, content_filler) in enumerate(sections):
        header_row = table.rows[sec_idx * 2]
        content_row = table.rows[sec_idx * 2 + 1]

        # ---- header row：合并所有列 + 灰底 + 左对齐粗体（对齐官方模板）----
        _set_row_height(header_row, header_h, exact=True)
        header_cell = _merge_row(header_row)
        _shade_cell(header_cell, HEADER_FILL)
        _vert_center(header_cell)
        _clear_and_fill_text(
            header_cell, label,
            size_pt=sec_pt, bold=True,
            align=WD_ALIGN_PARAGRAPH.LEFT,
        )

        # ---- content row：高度 + 由 filler 决定是否合并 ----
        _set_row_height(content_row, content_h)
        content_filler(content_row, outline)


# ----------------------------------------------------------------------------
# Section fillers
# ----------------------------------------------------------------------------
def _fill_difficulty(row, outline: BookOutline) -> None:
    cell = _merge_row(row)
    _vert_center(cell)
    _clear_paragraphs(cell)

    # Reader Type 按 Level 强制映射（v1.8）
    # - L0 (Smart) / L1 / L2 是固定文本
    # - L3-L6 = "Fiction" / "Non-Fiction"（取 outline.fiction_type 或 outline.reader_type）
    reader_type = _default_reader_type(outline)

    # 篇章难度（用户拍板 2026-06-09）：标签用「篇章难度」，值带 CEFR 前缀（如 "CEFR A2"）
    passage_cefr = _default_cefr_text(outline)

    lp = _lp(outline)
    diff_pt = lp.get("diff", DIFF_PT)
    ls = lp.get("ls_diff", lp.get("ls", LINE_SPACING))  # 放大行距铺满（短内容时）
    pa = lp.get("pa", PARA_AFTER_PT)

    pairs = [
        ("类型：", reader_type),
        ("阅读字数：", str(outline.total_words or "—")),
        ("篇章难度：", passage_cefr),
        ("语法难度：", lp.get("grammar") or _grammar_difficulty_text(outline) or "—"),
    ]
    for i, (label, value) in enumerate(pairs):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(pa)
        p.paragraph_format.line_spacing = ls
        r_lbl = p.add_run(label)
        _bind_run(r_lbl, FONT_EN, FONT_CN, size_pt=diff_pt, bold=True)
        r_val = p.add_run(value)
        _bind_run(r_val, FONT_EN, FONT_CN, size_pt=diff_pt, bold=False)
        if label == "语法难度：":
            # 语法难度的值显式取消加粗（否则会继承样式默认粗体）
            r_val.font.bold = False


def _fill_vocab(row, outline: BookOutline) -> None:
    """词汇掌握 —— 完全对齐官方模板（L5-1 5.27）：本行就是表格自身的 词数×2 个单元格，
    [词格] + [书写空格] 交替，紧凑单行（行高 ≈ 539），不嵌套子表、下方不留白。"""
    body_pt = _lp(outline).get("body", BODY_PT)
    words = [w for w in _vocab_words_for_rr(outline) if w] or [""]
    cells = row.cells
    for i, w in enumerate(words):
        wi = i * 2
        if wi >= len(cells):
            break
        wc = cells[wi]                 # 词格（居中显示目标词）
        _vert_center(wc)
        _clear_and_fill_text(
            wc, w, size_pt=body_pt, bold=False,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        if wi + 1 < len(cells):        # 书写空格（学生写中文意思）
            _vert_center(cells[wi + 1])


def _fill_phonics(row, outline: BookOutline) -> None:
    """自然拼读（L0-L4）/ 构词法（L5-L6）。"""
    cell = _merge_row(row)
    _vert_center(cell)
    _clear_paragraphs(cell)
    lp = _lp(outline)
    body_pt = lp.get("body", BODY_PT)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = lp.get("ls", LINE_SPACING)

    # 块6：命中大纲时拼读规则以大纲原文 verbatim 为准（防被 AI 推断覆盖）
    _phonics_src = outline.phonics
    _syl = getattr(outline, "syllabus", None)
    if _syl is not None and getattr(_syl, "phonics_rule", ""):
        _phonics_src = _syl.phonics_rule
    if _is_morphology_level(outline.level):
        text = _normalize_morphology(_phonics_src, outline)
    else:
        text = _enrich_phonics_examples(_normalize_phonics(_phonics_src), outline)
    text = _clean_rr_text(text)
    text = _concise_affix(text)   # 固定格式：去引号 + 保留例词 + 多条用 / 隔开
    run = p.add_run(text)
    _bind_run(run, FONT_EN, FONT_CN, size_pt=body_pt, bold=False)


def _fill_fluency(row, outline: BookOutline) -> None:
    cell = _merge_row(row)
    _vert_top(cell)   # 顶部对齐：短文从格子顶端开始，避免上下居中留白
    _clear_paragraphs(cell)

    lp = _lp(outline)
    ls = lp.get("ls_fluency", lp.get("ls", LINE_SPACING))  # 放大行距铺满正文区（短文时）
    body_text = _clean_rr_text(capitalize_names(_strip_wrapping_quotes(" ".join(
        _strip_wrapping_quotes(page.text.strip())
        for page in outline.pages
        if page.page_type == "story" and page.text and page.text.strip()
    ))))
    # 最终兜底：极端超长正文时按可容纳长度截断（确保严格 1 页，仅在降档仍溢出时触发）
    cap = getattr(outline, "_rr_passage_char_cap", None)
    if cap and len(body_text) > cap:
        body_text = body_text[:cap].rstrip().rstrip(",.;:") + "\u2026"
    # 字号由自适应规划器统一决定（保证 1 页）
    body_size = lp.get("body", BODY_PT)

    p_title = cell.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(0)
    p_title.paragraph_format.line_spacing = ls
    r_title = p_title.add_run(capitalize_names(outline.title or ""))
    _bind_run(r_title, FONT_EN, FONT_CN, size_pt=body_size, bold=False)

    if body_text:
        p_body = cell.add_paragraph()
        p_body.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_body.paragraph_format.space_before = Pt(0)
        p_body.paragraph_format.space_after = Pt(0)
        p_body.paragraph_format.line_spacing = ls
        r_body = p_body.add_run(body_text)
        _bind_run(r_body, FONT_EN, FONT_CN, size_pt=body_size, bold=False)


def _fill_questions(row, outline: BookOutline) -> None:
    cell = _merge_row(row)
    _vert_top(cell)   # 顶部对齐：题目在上、下方留白 = 学生答题书写空间
    _clear_paragraphs(cell)

    lp = _lp(outline)
    body_pt = lp.get("body", BODY_PT)
    ls = lp.get("ls_questions", lp.get("ls", LINE_SPACING))  # 放大行距铺满答题区（题少时）
    pa = lp.get("pa", PARA_AFTER_PT)

    questions = _resolve_rr_questions(outline)

    # 用户拍板 2026-06-08：L4-6 的阅读表达题不再标 (P#) 页码（⭐ 难度星保留）。L0-3 维持带页码。
    _lvl = (outline.level_key or "")
    _hide_page_ref = _lvl.isdigit() and int(_lvl) >= 4

    for i, q in enumerate(questions, start=1):
        p = cell.paragraphs[0] if i == 1 else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(pa)
        p.paragraph_format.line_spacing = ls

        # 题干：美式拼写 + 独立 i 大写 + 首字母大写 + 句末问号
        import re as _re
        text = _to_us_spelling(q["q"].strip().rstrip(".！？!?")).strip()
        text = capitalize_names(text)
        text = _re.sub(r"\bi\b", "I", text)
        text = _re.sub(r"\bi(['\u2019])", r"I\1", text)
        if text:
            text = text[0].upper() + text[1:]
        if text and text[-1] not in "?？.!":
            text += "?"
        page = q.get("page")
        stars = max(1, min(int(q.get("stars") or 1), 3))

        r_q = p.add_run(f"{i}. {text}")
        _bind_run(r_q, FONT_EN, FONT_CN, size_pt=body_pt, bold=False)

        if page is not None and not _hide_page_ref:
            r_p = p.add_run(f" (P{page})")
            _bind_run(r_p, FONT_EN, FONT_CN, size_pt=body_pt, bold=False)

        # 实心五角星 ★ (U+2605)，用阿里巴巴普惠体渲染（用户拍板：emoji ⭐ 在 WPS 显示为空心，
        # 改用 CJK 字体的实心 ★ + 橙色，保证填充实心）。
        r_s = p.add_run(" " + ("\u2605" * stars))
        _bind_run(r_s, FONT_CN, FONT_CN, size_pt=body_pt, bold=False)
        _set_run_color(r_s, STAR_COLOR)

        # 示例答案版：题干下方加一行灰色斜体示例答案（空白版不渲染）
        if getattr(outline, "_rr_with_answers", False):
            ans = (q.get("answer") or "").strip()
            if ans:
                ap = cell.add_paragraph()
                ap.alignment = WD_ALIGN_PARAGRAPH.LEFT
                ap.paragraph_format.space_before = Pt(0)
                ap.paragraph_format.space_after = Pt(pa)
                ap.paragraph_format.line_spacing = ls
                r_a = ap.add_run(f"   A: {capitalize_names(_to_us_spelling(ans))}")
                _bind_run(r_a, FONT_EN, FONT_CN, size_pt=body_pt, bold=False)
                r_a.italic = True
                _set_run_color(r_a, STAR_COLOR)


def _fill_engagement(row, outline: BookOutline) -> None:
    """课堂参与度：emoji + label + ☐ 复选框，3 组横向居中排列。"""
    cell = _merge_row(row)
    _vert_center(cell)
    _clear_paragraphs(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0

    # 对齐官方真人样本：笑脸 emoji + label + 空白方框（学生打勾）。
    # 😎 Excellent / 😄 Great / 🙂 Good，emoji 用 Segoe UI Emoji 渲染（彩色）。
    body_pt = _lp(outline).get("body", BODY_PT)
    # 统一三档笑脸 + label + 空格 + 方框（方框前留足空格，避免 emoji 与方框贴在一起/被挡）
    items = [("\U0001F606", "Excellent"), ("\U0001F604", "Great"), ("\U0001F642", "Good")]
    for idx, (face, label) in enumerate(items):
        r_f = p.add_run(face)
        _bind_run(r_f, EMOJI_FONT, EMOJI_FONT, size_pt=body_pt + 2, bold=False)
        r_l = p.add_run("  " + label + "   ")
        _bind_run(r_l, FONT_EN, FONT_CN, size_pt=body_pt, bold=False)
        r_box = p.add_run(EMPTY_BOX)
        _bind_run(r_box, SYMBOL_FONT, SYMBOL_FONT, size_pt=body_pt + 4, bold=False)
        if idx < len(items) - 1:
            r_gap = p.add_run("          ")
            _bind_run(r_gap, FONT_EN, FONT_CN, size_pt=body_pt, bold=False)


# ============================================================================
# 数据规整
# ============================================================================
def _rr_vocab_max(level: str) -> int:
    """RR 词汇格子数量上限（用户拍板 2026-06-04，分级 4-6 个）：
      • L0-2 / L3-4：4 个
      • L5-6：     最多 6 个（可含词组）
    """
    key = str(level or "").strip().lower()
    if "smart" in key:
        return 4
    digits = "".join(ch for ch in key if ch.isdigit())
    try:
        n = int(digits)
    except ValueError:
        return 4
    return 6 if n >= 5 else 4


def _vocab_words_for_rr(outline: BookOutline) -> list[str]:
    """RR 词汇掌握取词：按级别返回 4-6 个真实词（不补空字符串，便于格子自适应）。

    块6（用户拍板 2026-06-08）：命中大纲时，词形以大纲 verbatim 为准（防被 AI 抽取覆盖）。
    """
    syl = getattr(outline, "syllabus", None)
    if syl is not None:
        try:
            if outline.is_dual_vocab_level and getattr(syl, "vocab_mastery", None):
                syl_words = list(syl.vocab_mastery)
            else:
                syl_words = syl.vocab_words()
        except Exception:
            syl_words = []
        if syl_words:
            cleaned = [_to_us_spelling((w or "").strip().rstrip(",.;:").strip().lower())
                       for w in syl_words if (w or "").strip()]
            if cleaned:
                return cleaned[:_rr_vocab_max(outline.level)]
    if outline.is_dual_vocab_level and outline.vocabulary_mastery:
        words = list(outline.vocabulary_mastery)
    elif outline.vocabulary_simple:
        words = list(outline.vocabulary_simple)
    elif outline.vocabulary_mastery:
        words = list(outline.vocabulary_mastery)
    else:
        words = []
    cleaned = []
    for w in words:
        ww = (w or "").strip().rstrip(",.;:").strip()
        if ww:
            # v1.8：词汇统一小写、美式拼写
            cleaned.append(_to_us_spelling(ww.lower()))
    return cleaned[:_rr_vocab_max(outline.level)]


def _resolve_rr_questions(outline: BookOutline) -> list[dict]:
    """读取 outline._rr_questions（AI 抽取 / 人工编辑后挂载），按口径星级补齐。

    页码 (P#) 统一口径（所有级别一致，唯一规范）：
      • 事实定位题（⭐ / ⭐⭐）：必须带 (P#)，指向答案所在绘本页（P2-P8）。
      • 末尾开放拓展题（⭐⭐⭐，生活化/个人观点/PBL）：不带 (P#)（无法定位到具体某页）。

    若上游缺失，按 Level 题量给占位题。
    """
    raw = getattr(outline, "_rr_questions", None)
    dist = config.rr_question_distribution(outline.level)
    page_lookup = _story_page_lookup(outline)

    def _page_for(i: int, stars: int) -> int | None:
        # 末尾开放拓展题 ⭐⭐⭐ 无页码；事实题落在 P2-P8（i 从 0 起 → i+2）。
        return None if _rr_omit_page(stars) else (i + 2)

    if raw and isinstance(raw, list) and len(raw) > 0:
        normalized: list[dict] = []
        for i, q in enumerate(raw[:len(dist)]):
            raw_stars = q.get("stars") or q.get("difficulty") or q.get("level")
            stars = _parse_rr_stars(raw_stars) or dist[i]
            answer = str(q.get("answer") or q.get("sample") or "").strip()
            question = str(q.get("q") or q.get("question") or "").strip()
            located_page = _known_rr_page(outline, question, answer) or _page_from_question_answer(question, answer, page_lookup)
            page = located_page or (
                None if _rr_omit_page(stars) else (q.get("page") or (i + 2))
            )
            if not raw_stars and _literal_rr_question(question, answer):
                stars = 1
            normalized.append({
                "q": _normalize_student_punctuation(question),
                "stars": stars,
                "page": page,
                "answer": _normalize_student_punctuation(answer),
            })
        while len(normalized) < len(dist):
            i = len(normalized)
            stars = dist[i]
            normalized.append({
                "q": f"Question {i + 1}?",
                "stars": stars,
                "page": _page_for(i, stars),
            })
        return normalized

    return [
        {
            "q": f"Question {i + 1}?",
            "stars": stars,
            "page": _page_for(i, stars),
        }
        for i, stars in enumerate(dist)
    ]


def _rr_omit_page(stars: int) -> bool:
    """RR 页码统一规范（唯一事实源，已接入 _resolve_rr_questions）：
    末尾开放拓展题 ⭐⭐⭐ 不标 (P#)，其余事实题都标。"""
    try:
        return int(stars) >= 3
    except (TypeError, ValueError):
        return False


def _parse_rr_stars(value) -> int | None:
    if value is None:
        return None
    if isinstance(value, int):
        return max(1, min(value, 3))
    text = str(value).strip().lower()
    if not text:
        return None
    if "★★★" in text or "***" in text or "hard" in text or "challenge" in text:
        return 3
    if "★★" in text or "**" in text or "medium" in text:
        return 2
    if "★" in text or "*" in text or "easy" in text:
        return 1
    m = re.search(r"\b([123])\b", text)
    if m:
        return int(m.group(1))
    return None


def _story_page_lookup(outline: BookOutline) -> list[tuple[int, str]]:
    uploaded = getattr(outline, "_uploaded_book_page_lookup", None) or []
    clean_uploaded: list[tuple[int, str]] = []
    for page_no, text in uploaded:
        try:
            no = int(page_no)
        except (TypeError, ValueError):
            continue
        body = (text or "").strip().lower()
        if no > 0 and body:
            clean_uploaded.append((no, body))
    if clean_uploaded:
        return clean_uploaded

    pages: list[tuple[int, str]] = []
    sentence_texts: list[str] = []
    for page in outline.pages:
        text = (getattr(page, "text", "") or "").strip()
        if getattr(page, "page_type", "") == "story" and text:
            pages.append((page.index + 1, text.lower()))
            sentence_texts.extend(
                s.strip().lower()
                for s in re.split(r"(?<=[.!?])\s+", text)
                if s.strip()
            )
    if len(sentence_texts) >= 8:
        sent_pages: list[tuple[int, str]] = []
        denom = max(1, len(sentence_texts) - 1)
        for i, sent in enumerate(sentence_texts):
            printed_page = int(round(2 + i * 6 / denom))
            printed_page = max(2, min(8, printed_page))
            sent_pages.append((printed_page, sent))
        # Prefer sentence-level lookup; keep original page chunks as fallback.
        return sent_pages + pages
    return pages


def _page_from_question_answer(question: str, answer: str, page_lookup: list[tuple[int, str]]) -> int | None:
    q = (question or "").strip().lower()
    ans = (answer or "").strip().lower()
    if not (q or ans):
        return None
    for page_no, text in page_lookup:
        if ans and ans in text and len(ans) >= 12:
            return page_no
    stop = {
        "what", "when", "where", "which", "will", "does", "did", "the", "and",
        "this", "that", "with", "many", "much", "next",
    }
    words = [
        w for w in re.findall(r"[a-zA-Z]+", f"{q} {ans}")
        if len(w) >= 4 and w not in stop
    ]
    if not words:
        return None
    best_page = None
    best_score = 0
    for page_no, text in page_lookup:
        score = sum(1 for w in words if w.lower() in text)
        if score > best_score:
            best_page = page_no
            best_score = score
    return best_page if best_score >= min(2, len(words)) else None


def _known_rr_page(outline: BookOutline, question: str, answer: str = "") -> int | None:
    """Book-specific guardrails for exact RR page refs when the uploaded OCR split is noisy."""
    title = (getattr(outline, "title", "") or "").lower()
    if "mia" in title and "seven-day plan" in title:
        text = f"{question} {answer}".lower()
        mapping = [
            (("room", "messy"), 3),
            (("show", "sunday"), 4),
            (("do", "first"), 6),
            (("practice", "every", "day"), 7),
            (("happy",), 8),
        ]
        for keys, page in mapping:
            if all(k in text for k in keys):
                return page
    return None


def _literal_rr_question(question: str, answer: str = "") -> bool:
    text = f"{question} {answer}".lower()
    if any(k in text for k in ("why", "how do you", "what do you think", "your")):
        return False
    literal_keys = ("what", "when", "where", "which", "is ", "does ", "will ")
    return any(k in f" {text}" for k in literal_keys)


def _grammar_difficulty_text(outline: BookOutline) -> str:
    """RR grammar difficulty should name the tense, not only the sentence frame."""
    story = " ".join(
        (p.text or "").strip()
        for p in getattr(outline, "pages", []) or []
        if getattr(p, "page_type", "") == "story" and (p.text or "").strip()
    ).lower()
    found: list[str] = []
    if re.search(r"\b(am|is|are|has|have|do|does|want|wants|feel|feels|make|makes)\b", story):
        found.append("一般现在时态")
    if re.search(r"\bwill\s+[a-z]+\b", story):
        found.append("一般将来时态")
    if re.search(r"\b(was|were|had|did|went|saw|felt|made|played|wanted|practiced|cleaned)\b", story):
        found.append("一般过去时态")
    if "一般将来时态" in found:
        simple = [x for x in dict.fromkeys(found) if x != "一般将来时态"]
        simple.append("will + 动词原形")
        return "；".join(simple)
    if found:
        return "、".join(dict.fromkeys(found))
    return _normalize_grammar_cn(getattr(outline, "grammar_focus", "")) or ""


def _clean_rr_text(text: str) -> str:
    return _normalize_student_punctuation(text)


def _default_reader_type(outline) -> str:
    """按 Level 强制映射 Reader Type（v1.8 规范）。

      Smart / L0 → Concept & Knowledge - Building Readers
      L1         → Patterned Narrative & Informational Readers
      L2         → Early Independent Genre-Exposure Readers
      L3 - L6    → Fiction  或  Non-Fiction   （取 outline.fiction_type / reader_type）

    用户在大纲中显式写了 reader_type 时，对 L0-L2 仍按级别覆盖；对 L3-L6 优先用 reader_type。
    """
    key = str(getattr(outline, "level", "") or "").strip().lower()
    if "smart" in key:
        return "Concept & Knowledge - Building Readers"
    digits = "".join(ch for ch in key if ch.isdigit())
    try:
        n = int(digits)
    except ValueError:
        return "Patterned Narrative & Informational Readers"
    if n == 0:
        return "Concept & Knowledge - Building Readers"
    if n == 1:
        return "Patterned Narrative & Informational Readers"
    if n == 2:
        return "Early Independent Genre-Exposure Readers"
    # L3 - L6: Fiction / Non-Fiction
    ft = (getattr(outline, "fiction_type", "") or getattr(outline, "reader_type", "") or "").strip()
    ft_low = ft.lower()
    if "non" in ft_low and "fic" in ft_low:
        return "Non-Fiction"
    if "fic" in ft_low:
        return "Fiction"
    return "Fiction"  # L3-L6 没填默认 Fiction


def _default_cefr_text(outline) -> str:
    """词汇难度长文本（带 CEFR 前缀）— 保留兼容旧调用。"""
    if getattr(outline, "cefr", "") and outline.cefr.strip():
        return f"CEFR {outline.cefr.strip()}"
    return f"CEFR {_default_cefr_short_code(outline)}"


def _default_cefr_short_code(outline) -> str:
    """词汇难度短码（按官方 L5-1 实测格式，无 CEFR 前缀，无 Lexile）：

      Smart → Pre-A1
      L0    → Pre-A1
      L1    → Pre-A1
      L2    → A1
      L3    → A1+
      L4    → A2
      L5    → A2     # 官方 L5-1 实测是 A2（不是 B1，官方更保守）
      L6    → B1

    若用户在大纲显式写了 cefr 字段，优先用大纲值。
    """
    if getattr(outline, "cefr", "") and outline.cefr.strip():
        # 移除可能的 "CEFR " 前缀
        cefr = outline.cefr.strip()
        if cefr.upper().startswith("CEFR "):
            cefr = cefr[5:].strip()
        return cefr
    key = str(getattr(outline, "level", "") or "").strip().lower()
    if "smart" in key:
        return "Pre-A1"
    digits = "".join(ch for ch in key if ch.isdigit())
    mapping = {
        0: "Pre-A1", 1: "Pre-A1", 2: "A1", 3: "A1+",
        4: "A2",     5: "A2",     6: "B1",
    }
    try:
        return mapping.get(int(digits), "A1")
    except ValueError:
        return "A1"


def _normalize_grammar_cn(raw: str) -> str:
    """语法难度统一显示中文时态名（v1.8.3 规则）。

    检测英文时态短语关键词 → 中文标准名：
      Simple present tense / present simple → 一般现在时态
      Simple past tense / past simple       → 一般过去时态
      Simple future tense                   → 一般将来时态
      Present continuous / present progressive → 现在进行时态
      Past continuous                       → 过去进行时态
      Present perfect                       → 现在完成时态
      Past perfect                          → 过去完成时态
      Modal verbs                            → 情态动词

    如果包含其他句型（"There is/are", "help + V" 等），按 "时态 + 句型" 拼接。
    用户已经填了中文则直接返回。
    """
    import re as _re
    s = (raw or "").strip()
    if not s:
        return ""
    # 已经含中文时态字眼 → 去掉冗长英文括注（如动词清单）后返回，保留简短中文括注
    if _re.search(r"[一二三过将完现][般去来成在]", s):
        def _drop_long_paren(m):
            inner = m.group(1)
            # 长括注（>14 字）且含英文字母 → 删除（避免语法行撑成多行）
            if len(inner) > 14 and _re.search(r"[A-Za-z]", inner):
                return ""
            return m.group(0)
        s = _re.sub(r"\s*[\(（]([^)）]*)[\)）]", _drop_long_paren, s)
        return s.rstrip("。.；; ").strip()

    low = s.lower()
    tense_map = [
        ("present perfect", "现在完成时态"),
        ("past perfect", "过去完成时态"),
        ("present continuous", "现在进行时态"),
        ("present progressive", "现在进行时态"),
        ("past continuous", "过去进行时态"),
        ("past progressive", "过去进行时态"),
        ("simple future", "一般将来时态"),
        ("future simple", "一般将来时态"),
        ("simple past", "一般过去时态"),
        ("past simple", "一般过去时态"),
        ("past tense", "一般过去时态"),
        ("simple present", "一般现在时态"),
        ("present simple", "一般现在时态"),
        ("present tense", "一般现在时态"),
        ("modal verb", "情态动词"),
        ("imperative", "祈使句"),
    ]
    parts: list[str] = []
    found_tense = False
    for en, cn in tense_map:
        if en in low and cn not in parts:
            parts.append(cn)
            found_tense = True
            break
    # 句型附加
    if "there was" in low or "there were" in low:
        parts.append("there was/were 句型")
    elif "there is" in low or "there are" in low:
        parts.append("there is/are 句型")
    if "+ v" in low or "+ verb" in low:
        # 取动词搭配描述（如 "help + V" → "help + V 句型"）
        m = _re.search(r"(\w+)\s*\+\s*v", low)
        if m:
            parts.append(f"{m.group(1)} + V 句型")
    if not found_tense and not parts:
        # 没匹配任何时态 → 原文保留
        return s.rstrip("。.；; ").strip()
    return "；".join(parts)


def _level_number(level: str) -> int:
    """从 level 字符串解析数字级别；解析不出返回 -1（如 smart）。"""
    key = str(level or "").strip().lower()
    if "smart" in key:
        return -1
    digits = "".join(ch for ch in key if ch.isdigit())
    try:
        return int(digits)
    except ValueError:
        return -1


def _is_high_level(level: str) -> bool:
    """L3-6 视为高级别（页首标题用四号 14pt）。"""
    return _level_number(level) >= 3


def _is_morphology_level(level: str) -> bool:
    """L5 / L6 用构词法 (morphology)；其他级别用自然拼读 (phonics)。"""
    key = str(level or "").strip().lower()
    if "smart" in key:
        return False
    digits = "".join(ch for ch in key if ch.isdigit())
    try:
        n = int(digits)
    except ValueError:
        return False
    return n >= 5


# 常见后缀 / 前缀 → 含义（构词法兜底库，从故事词汇里自动检测）
_SUFFIX_MEANINGS = {
    "-ous":  "having/full of a quality",
    "-ful":  "full of",
    "-less": "without",
    "-able": "can be / capable of",
    "-tion": "act/state of",
    "-ment": "result/action of",
    "-ness": "state/quality of",
    "-ly":   "in a ... manner (adverb)",
    "-er":   "person who / more",
    "-est":  "most (superlative)",
    "-ed":   "past tense / past participle",
    "-ing":  "continuous / gerund",
    "-y":    "having / characterized by",
    "-ish":  "somewhat / having quality of",
    "-en":   "to make / become",
    "-ity":  "quality / condition of",
}
_PREFIX_MEANINGS = {
    "un-":  "not / opposite of",
    "re-":  "again",
    "pre-": "before",
    "dis-": "not / opposite of",
    "in-":  "not",
    "im-":  "not",
    "non-": "not",
    "mis-": "wrongly / badly",
    "over-": "too much",
    "under-": "too little / beneath",
}


def _detect_morphology_in_words(words: list[str]) -> Optional[str]:
    """从词汇里检测最高频的构词法规则（后缀优先于前缀）。

    返回固定格式（用户拍板 2026-06-04）：
      'suffix -ous (= having/full of a quality): nervous, famous'
      —— 去引号、冒号分隔例词、**最多 2 例**、保证一行。
    """
    if not words:
        return None
    words = [str(w or "").strip().lower() for w in words if str(w or "").strip()]

    # 后缀检测：词尾匹配
    for suf, meaning in _SUFFIX_MEANINGS.items():
        suf_letters = suf.lstrip("-")
        hits = [w for w in words if len(w) > len(suf_letters) + 1 and w.endswith(suf_letters)]
        if len(hits) >= 1:
            examples = ", ".join(hits[:2])
            return f"suffix {suf} (= {meaning}): {examples}"

    # 前缀检测：词头匹配
    for pre, meaning in _PREFIX_MEANINGS.items():
        pre_letters = pre.rstrip("-")
        hits = [w for w in words if len(w) > len(pre_letters) + 1 and w.startswith(pre_letters)
                # 排除 "income" 误判 "in-" 等：前缀后接元音/辅音规则太复杂，简化为只在词长>=4 时认
                and len(w) >= 4]
        if len(hits) >= 1:
            examples = ", ".join(hits[:2])
            return f"prefix {pre} (= {meaning}): {examples}"

    return None


def _concise_affix(text: str) -> str:
    """构词法/自然拼读固定格式（用户拍板 2026-06-04）：

      • 自然拼读：'short a (cat, hat, map)'  —— 音 + 例词括注
      • 构词法：  'suffix -ous (= having/full of a quality): nervous, famous'
                  —— 规则(含义) + 冒号分隔例词、**最多 2 例**、一行
      • 多条规则保留，用 ' / ' 隔开，继续相同模式
      • 一律去掉词缀/音素外层引号（"-ous" → -ous，"a" → a）
    """
    import re as _re
    # 只在分号或"空格 / 空格"处断成多条规则；含义里的裸 "/"（having/full）不拆
    rules = [r.strip() for r in _re.split(r"\s*[;；]\s*|\s+/\s+", (text or "").strip()) if r.strip()]
    out: list[str] = []
    for r in rules:
        # 注意：自然拼读(phonics)需保留英文直双引号（如 long vowel "ea" as in "health"），
        # 只在构词法分支里去掉词缀外层引号；这里仅去掉单引号。
        r = r.replace("'", "").strip().rstrip("；;").strip()
        if not r:
            continue

        # 构词法分支：含 "(= 含义)" → 规整成 'head (= 含义): ex1, ex2'（最多 2 例）
        m_mean = _re.search(r"\(=\s*(.*?)\)", r)
        if m_mean:
            head = r[:m_mean.start()].replace('"', "").strip()
            meaning = m_mean.group(1).strip()
            rest = r[m_mean.end():].strip()

            examples: list[str] = []
            # 含义里内嵌的 "..., e.g. ex" 抽出来
            me = _re.search(r",?\s*e\.g\.\s*(.+)$", meaning)
            if me:
                examples += [e.strip() for e in _re.split(r"[,，]", me.group(1)) if e.strip()]
                meaning = meaning[:me.start()].strip().rstrip(",").strip()
            # 尾部例词：可能是 ": ex1, ex2" / "(ex1, ex2)" / ", e.g. ex"
            rest = rest.lstrip(":：").strip().strip("()").strip()
            rest = _re.sub(r"^e\.g\.\s*", "", rest)
            examples += [e.strip() for e in _re.split(r"[,，]", rest) if e.strip()]
            examples = [e for e in examples if e][:2]

            res = f"{head} (= {meaning})" if meaning else head
            if examples:
                res += ": " + ", ".join(examples)
            out.append(res)
            continue

        # 自然拼读分支：把 "音: 例词" 冒号写法转成括注： short a: cat, hat → short a (cat, hat)
        m = _re.search(r"^(.*?\S)\s*[:：]\s*([A-Za-z][A-Za-z,\s/]*)$", r)
        if m and "(" not in r:
            r = f"{m.group(1).strip()} ({m.group(2).strip()})"
        out.append(r)
    return "  /  ".join(out)


def _normalize_morphology(raw: str, outline) -> str:
    """L5-L6 构词法格式化。

    - 用户填了 phonics 字段就用它（仍按 _normalize_phonics 规整 quote / 大小写）
    - 没填则从 outline 的 mastery / vocabulary 词里自动检测后缀/前缀
    - 都失败兜底为 'suffix "-ly" (= in a ... manner)'
    """
    if raw and str(raw).strip():
        return _normalize_phonics(raw)

    words: list[str] = []
    if getattr(outline, "vocabulary_mastery", None):
        words.extend(outline.vocabulary_mastery)
    if getattr(outline, "vocabulary_simple", None):
        words.extend(outline.vocabulary_simple)
    if getattr(outline, "vocabulary_exposure", None):
        words.extend(outline.vocabulary_exposure)

    auto = _detect_morphology_in_words(words)
    if auto:
        return auto
    return 'suffix "-ly" (= in a ... manner, adverb)'


def _normalize_phonics(raw: str) -> str:
    """规整 phonics 文本为 v1.8 sample 风格（不是句子，是词组）：

    规则：
      • 全小写（CEFR / PBL 等术语缩写保持大写）
      • 英文直双引号 \"...\"（自动把 curly quote 转回 straight quote）
      • 例词放在括号里：(friendship)
      • 不带句号
      • 多条规则用分号 + 空格分隔

    示例输入 → 输出：
      'Consonant blend FR (friendship).' → 'consonant blend \"fr\" (friendship)'
      'long o (snow ow)'                  → 'long \"o\" (snow ow)'
      'AI → /eɪ/: day, stay, play'        → 'ai → /eɪ/ (day, stay, play)'
      ''                                   → 'short vowel pattern'
    """
    import re

    if not raw:
        return "short vowel pattern"
    text = _normalize_student_punctuation(str(raw).strip().replace("\r\n", " ").replace("\n", " "))

    # 1. curly quote → straight quote
    text = (
        text.replace("\u201c", '"').replace("\u201d", '"')
            .replace("\u2018", '"').replace("\u2019", '"')
    )
    # 2. 去掉末尾句号
    text = text.rstrip(". ").strip()

    # 3. 整体小写化（但 CEFR / PBL 大写术语保留）— 直接全 lower，再把术语还原
    preserved = ["CEFR", "PBL"]
    text = text.lower()
    for term in preserved:
        text = re.sub(rf"\b{term.lower()}\b", term, text)

    # 4. ": word, word" → " (word, word)"
    m = re.search(r":\s*([a-zA-Z][a-zA-Z,\s]*[a-zA-Z])\s*$", text)
    if m:
        words = m.group(1).strip()
        text = text[:m.start()].rstrip() + f" ({words})"

    # 5. 若文本里完全没有双引号，尝试给"音素词"加双引号
    #    模式 a: <word(s)> <single token without quotes> ( <example words> )
    #    例如 "consonant blend fr (friendship)" → consonant blend "fr" (friendship)
    if '"' not in text:
        m2 = re.search(r"^(.*\S)\s+(\S+)\s*\(([^)]+)\)\s*$", text)
        if m2:
            head, token, examples = m2.group(1), m2.group(2), m2.group(3)
            # 防止把已经是 "blend"/"vowel" 等普通修饰词加引号 — 用启发式：
            # token 仅 1-4 个字母且不在常见 stopwords 里
            stop = {"blend", "vowel", "consonant", "long", "short", "digraph",
                    "diphthong", "the", "and"}
            if token.lower() not in stop and 1 <= len(token) <= 5 and token.replace("+", "").isalpha():
                text = f'{head} "{token}" ({examples})'
    return text


def _enrich_phonics_examples(text: str, outline: BookOutline) -> str:
    low = (text or "").lower()
    if "ay" not in low:
        return text
    story = " ".join(
        (getattr(p, "text", "") or "")
        for p in getattr(outline, "pages", []) or []
        if getattr(p, "page_type", "") == "story"
    )
    examples: list[str] = []
    for word in re.findall(r"\b[A-Za-z]*ay[A-Za-z]*\b", story):
        clean = word.strip(".,!?;:").lower()
        if clean and clean not in examples:
            examples.append(clean)
    priority = [w for w in ("day", "play", "sunday", "tuesday") if w in examples]
    rest = [w for w in examples if w not in priority]
    chosen = (priority + rest)[:5]
    if not chosen:
        return text
    chosen = [w.capitalize() if w in {"sunday", "tuesday"} else w for w in chosen]
    base = re.sub(r"\s*\([^)]*\)\s*$", "", text).strip()
    return f"{base} ({', '.join(chosen)})"


def _level_label(level: str) -> str:
    """Level 0 / Smart / 1-6 → 标题里的 'Level X' 或 'Smart'."""
    s = str(level or "").strip()
    low = s.lower()
    if low.startswith("smart"):
        return "Smart"
    digits = "".join(ch for ch in s if ch.isdigit())
    return f"Level {digits or '1'}"


# ============================================================================
# python-docx XML 工具
# ============================================================================
def _bind_run(run, ascii_font: str, east_asia: str, size_pt: int, bold: bool) -> None:
    """每个 run 必须双绑 ascii + eastAsia 字体，强制规整。"""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), ascii_font)
    rFonts.set(qn("w:hAnsi"), ascii_font)
    rFonts.set(qn("w:eastAsia"), east_asia)
    rFonts.set(qn("w:cs"), ascii_font)

    half_pt = str(int(size_pt * 2))
    sz = rPr.find(qn("w:sz"))
    if sz is None:
        sz = OxmlElement("w:sz")
        rPr.append(sz)
    sz.set(qn("w:val"), half_pt)
    szCs = rPr.find(qn("w:szCs"))
    if szCs is None:
        szCs = OxmlElement("w:szCs")
        rPr.append(szCs)
    szCs.set(qn("w:val"), half_pt)

    if bold:
        if rPr.find(qn("w:b")) is None:
            rPr.append(OxmlElement("w:b"))
        if rPr.find(qn("w:bCs")) is None:
            rPr.append(OxmlElement("w:bCs"))


def _set_default_style(doc) -> None:
    style = doc.styles["Normal"]
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), FONT_EN)
    rFonts.set(qn("w:hAnsi"), FONT_EN)
    rFonts.set(qn("w:eastAsia"), FONT_CN)
    rFonts.set(qn("w:cs"), FONT_EN)


def _set_a4_portrait(doc) -> None:
    """A4 portrait 8.27×11.69 in，margins 1.0 cm（紧凑版式确保 1 页）"""
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(1.0)
        section.right_margin = Cm(1.0)
        section.top_margin = Cm(0.9)
        section.bottom_margin = Cm(0.9)


def _set_tbl_w(table, width_dxa: int) -> None:
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(width_dxa))
    tblW.set(qn("w:type"), "dxa")


def _set_tbl_borders(table) -> None:
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for kind in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{kind}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "BFBFBF")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _set_tbl_no_borders(table) -> None:
    """把表格所有边框设为 none（用于标题行的无边框布局表）。"""
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    tblBorders = OxmlElement("w:tblBorders")
    for kind in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{kind}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        b.set(qn("w:space"), "0")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _set_tbl_grid(table, widths_dxa: list[int]) -> None:
    tbl = table._element
    grid = tbl.find(qn("w:tblGrid"))
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is not None:
            tblPr.addnext(grid)
        else:
            tbl.insert(0, grid)
    for gc in list(grid.findall(qn("w:gridCol"))):
        grid.remove(gc)
    for w in widths_dxa:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        grid.append(gc)


def _set_row_height(row, twips: int, *, exact: bool = False) -> None:
    trPr = row._tr.get_or_add_trPr()
    h = trPr.find(qn("w:trHeight"))
    if h is None:
        h = OxmlElement("w:trHeight")
        trPr.append(h)
    h.set(qn("w:val"), str(twips))
    h.set(qn("w:hRule"), "exact" if exact else "atLeast")


def _merge_row(row):
    cells = row.cells
    merged = cells[0]
    for c in cells[1:]:
        merged = merged.merge(c)
    return merged


def _shade_cell(cell, hex_fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)


def _vert_center(cell) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _vert_top(cell) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def _set_run_color(run, hex_color: str) -> None:
    rPr = run._element.get_or_add_rPr()
    color = rPr.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        rPr.append(color)
    color.set(qn("w:val"), hex_color)


def _clear_paragraphs(cell) -> None:
    paras = cell.paragraphs
    if not paras:
        return
    first = paras[0]
    for run in list(first.runs):
        first._element.remove(run._element)
    for p in paras[1:]:
        p._element.getparent().remove(p._element)


def _clear_and_fill_text(cell, text: str, *, size_pt: int = 14,
                         bold: bool = False,
                         align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    _clear_paragraphs(cell)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    _bind_run(run, FONT_EN, FONT_CN, size_pt=size_pt, bold=bold)


# ============================================================================
# 兼容入口
# ============================================================================
def attach_rr_questions(outline: BookOutline, questions: list[dict]) -> None:
    setattr(outline, "_rr_questions", questions)
