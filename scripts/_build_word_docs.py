# -*- coding: utf-8 -*-
"""生成两份 Word：① 30 分钟汇报提词卡；② 绘本 4 件套详细介绍（表格）。
运行：.venv\\Scripts\\python.exe scripts\\_build_word_docs.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(r"C:\Users\Jered\picturebook-auto\outputs")
CN = "Microsoft YaHei"
ORANGE = "F76B1C"
ORANGE_RGB = RGBColor(0xF7, 0x6B, 0x1C)
INK = RGBColor(0x23, 0x27, 0x2E)
GRAY = RGBColor(0x6B, 0x72, 0x80)


def set_cjk(run):
    run.font.name = CN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CN)


def base_doc():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = CN
    st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), CN)
    for sec in doc.sections:
        sec.top_margin = Inches(0.7)
        sec.bottom_margin = Inches(0.7)
        sec.left_margin = Inches(0.7)
        sec.right_margin = Inches(0.7)
    return doc


def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hex_color)
    tcPr.append(sh)


def cell_text(cell, text, *, bold=False, color=None, size=10, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    parts = text.split("\n")
    for i, line in enumerate(parts):
        if i:
            p.add_run().add_break()
        r = p.add_run(line)
        r.bold = bold
        r.font.size = Pt(size)
        if color is not None:
            r.font.color.rgb = color
        set_cjk(r)


def title(doc, text, sub=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = ORANGE_RGB
    set_cjk(r)
    if sub:
        ps = doc.add_paragraph()
        rs = ps.add_run(sub)
        rs.font.size = Pt(10.5)
        rs.font.color.rgb = GRAY
        set_cjk(rs)
    bar = doc.add_paragraph()
    bp = bar.paragraph_format
    bp.space_before = Pt(2)
    bp.space_after = Pt(6)


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = INK
    set_cjk(r)


def para(doc, text, *, color=INK, size=10.5, bold=False, space=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color
    set_cjk(r)
    return p


def make_table(doc, headers, rows, widths=None, header_fill=ORANGE):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        shade(hdr[i], header_fill)
        cell_text(hdr[i], h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=10.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cell_text(cells[i], str(val), size=10)
    if widths:
        for col, w in enumerate(widths):
            for cell in t.columns[col].cells:
                cell.width = Inches(w)
    return t


# ============================================================
# 文档 1：30 分钟汇报提词卡
# ============================================================
def build_runsheet():
    doc = base_doc()
    title(doc, "绘本 4 件套自动化 · 30 分钟汇报提词卡",
          "面向 同事 + 领导（领导优先）｜ Hook → SCQA → 现场 Demo（高潮）→ 诚实复盘 → 目标 + 邀请试用")

    h2(doc, "一、流程编排（Run of Show）")
    make_table(
        doc,
        ["时段", "模块", "时长", "PPT 页", "一句话要点"],
        [
            ["0–2′", "开场 Hook", "2′", "P1 封面", "「大半天 vs 3 分钟」反差，敢现场赌一把"],
            ["2–7′", "S 背景 + C 冲突", "5′", "P3 · P4", "人工重 / 修图费劲 / 出题不匹配 / 格式因人而异；时间紧又要保质保量"],
            ["7–9′", "Q 提问 + A 回答", "2′", "P5", "把规范 / IP / 画风 / 题型沉淀进网页工具"],
            ["9–14′", "方案 + 三大能力", "5′", "P6 · P7", "锁 IP / 统一画风(gpt-image-2) / 规范落地"],
            ["14–22′", "现场 Demo（高潮）", "8′", "切到网页", "贴原文 → AI 抽取 → 微调一处 → 一键 4 件套"],
            ["22–24′", "Before / After", "2′", "P8", "标准从「人」搬到「系统」"],
            ["24–26′", "价值 + 诚实复盘", "2′", "P9", "提效 / 标准 / 规模；坦诚未解决项"],
            ["26–28′", "下一步目标", "2′", "P10", "正确率 ≥85% · 稳定率 ≥85% · 批量测试"],
            ["28–30′", "结尾", "2′", "P11", "给大家用 · 欢迎反馈，金句收尾"],
        ],
        widths=[0.8, 1.7, 0.6, 1.0, 3.4],
    )

    h2(doc, "二、关键话术（可直接念）")
    para(doc, "▍开场 Hook（0–2′）", color=ORANGE_RGB, bold=True, space=2)
    para(doc, "做一本绘本全套 4 件套，IC 配合教研复查要大半天到一天，还常返工改图。今天我想用约 3 分钟，"
              "当着大家的面，从一段故事原文现场生成这一整套——若能成，0–6 级绘本第一次能规模化、标准统一地铺出来。")
    para(doc, "▍冲突金句（C · 戳痛点）", color=ORANGE_RGB, bold=True, space=2)
    para(doc, "问题不在某个人不努力，而在于质量挂在「人」身上——人一换、活一多，标准就守不住。")
    para(doc, "▍一句话答案（A）", color=ORANGE_RGB, bold=True, space=2)
    para(doc, "把规范 / IP / 画风 / 题型全部沉淀进一个网页工具——一段原文进去，AI 抽取、老师微调、一键出 4 件套。")
    para(doc, "▍诚实复盘（24–26′ · 最赢信任）", color=ORANGE_RGB, bold=True, space=2)
    para(doc, "还没完全解决：① 高级别复杂画面偶有偏差需人工微调；② 文本正确率还在冲 90%+；③ 批量稳定性待大规模验证。"
              "收口：「它现在是好用的助手，还不是全自动，但方向已经跑通。」")
    para(doc, "▍结尾金句（28–30′）", color=ORANGE_RGB, bold=True, space=2)
    para(doc, "让 IC 从「人肉生图工具」回到「内容判断者」，让标准长在系统里、不长在某个人身上。"
              "工具已开放给大家试用，欢迎随时反馈，我会持续打磨。")

    h2(doc, "三、Demo 防翻车清单（高潮 8′）")
    for line in [
        "开场前先完整跑通一次，确认能出 4 件套",
        "留一份已生成好的成品做兜底",
        "确认网页服务已启动（localhost:8501）、网络可用、浏览器已开到首页",
        "边演边说：「我没写代码、没调参数，只是确认和微调」",
        "网络 / 接口慢 → 立刻切成品继续讲，节奏不停",
    ]:
        para(doc, "□  " + line, space=2)

    h2(doc, "四、对象共鸣")
    make_table(
        doc,
        ["对象", "他们关心什么 / 你怎么打"],
        [
            ["领导", "值不值得做（效率/ROI）· 现在到什么程度（可信、不吹）· 赋能不是替代"],
            ["同事", "会不会抢活 · 好不好用 · 帮我省哪些苦活 → 从「画图工」回到「内容判断者」"],
        ],
        widths=[1.0, 5.5],
    )

    dest = OUT / "汇报提词卡_30分钟.docx"
    doc.save(dest)
    return dest


# ============================================================
# 文档 2：绘本 4 件套详细介绍
# ============================================================
def build_suite_doc():
    doc = base_doc()
    title(doc, "绘本课程「4 件套」· 详细介绍",
          "一段故事原文 → AI 抽取 → 老师微调 → 一键产出 4 件「连续、互相对得上」的教学物料")

    h2(doc, "总览")
    make_table(
        doc,
        ["交付物", "格式", "篇幅", "一句话定位"],
        [
            ["① Picture Book 绘本", ".pptx", "4 的倍数页（1 封面 + 7 故事 + 结尾）", "孩子主读物：7 句故事 + 整本插画"],
            ["② Worksheet 练习作业纸", ".pptx", "6 页", "随书练习：5 类题型巩固词汇/句型/阅读"],
            ["③ Reading Report 阅读报告", ".docx", "1 页", "单页能力卡：阅读量/词汇/拼读·构词"],
            ["④ Teacher's Guide 教师指南", ".docx", "多页（套官方母版）", "授课指引：目标/流程/活动/引导问"],
        ],
        widths=[1.9, 0.7, 2.3, 2.4],
    )

    # 1 Picture Book
    h2(doc, "① Picture Book（绘本 PPT）")
    make_table(
        doc,
        ["维度", "标准 / 说明"],
        [
            ["定位", "整套核心读物，孩子主读物：7 句故事 + 整本插画"],
            ["页数结构", "总页数必须为 4 的倍数；1 封面 + 7 故事页 + 1 结尾/封底（不足用空白页凑齐 4 的倍数）"],
            ["封面", "顶部预留 Logo 位，书名标题置于页面上方，构图上空留白"],
            ["页码", "正文从 P2 开始；页码左下 / 右下交替，位置固定不变"],
            ["封底信息（6 项必齐）", "Level / Book number / CEFR / Lexile / Word count / Vocabulary"],
            ["词汇表", "分两行：第 1 行 Mastery 掌握词，第 2 行 Exposure 认知词"],
            ["插画生成", "gpt-image-2 生成，3:2 横版（1536×1024）；每页预留 10–15% 文字留白，避免台词遮挡主体"],
            ["画风", "温暖治愈水彩童书风，低饱和、柔和晕染、线条圆润；无畸形肢体、无过曝；跨页/册统一"],
            ["IP 一致性", "参考图锁定固定 IP 家族（Mia/Tommy 等）；按级别配年龄（L0–3=8 岁 / L4–5=10 岁 / L6=12 岁）；每页只改表情动作"],
            ["字体", "英文 Poppins：封面标题 Bold ≈40；内页正文 20–24，行高 1.2–1.5，不遮挡配图"],
        ],
        widths=[1.9, 4.6],
    )

    # 2 Worksheet
    h2(doc, "② Worksheet（练习作业纸 PPT）")
    make_table(
        doc,
        ["维度", "标准 / 说明"],
        [
            ["定位", "随书练习纸，巩固词汇 / 句型 / 阅读理解"],
            ["篇幅", "6 页（对齐 VIPKID 真实样本）"],
            ["题型（5 类）", "① Vocabulary 词义配对　② 单词补全拼写　③ 单选勾选　④ 动词过去式改写造句　⑤ 阅读单选"],
            ["难度梯度", "题型由易到难排序，靠后/下方题目按需更简单；一题只问一个明确问题"],
            ["版式", "粉色通栏标题条 + 左上角 Dino logo + 右上角 Name 标签；按单题文字容量自适应框大小，仅同列基准对齐"],
            ["字号分级（4 档）", "大标题 > 题型说明 > 题干正文 > Hint 小字；大标题 Poppins Bold 40 居中黑色，题干 20pt、说明 12pt"],
            ["配图", "尽量配图，帮孩子明确要写/输出的词是什么"],
            ["文本格式", "美式拼写；单词小写、不加句号；句子首字母大写 + 句号；活动须落到明确输出（color 只是手段不算 reading）"],
        ],
        widths=[1.9, 4.6],
    )

    # 3 Reading Report
    h2(doc, "③ Reading Report（阅读报告 DOCX）")
    make_table(
        doc,
        ["维度", "标准 / 说明"],
        [
            ["定位 / 篇幅", "单页阅读能力卡，所有内容必须压在 1 页内、不溢出"],
            ["阅读类型（首行，按级别固定）",
             "L0(Smart)：Concept & Knowledge-Building Readers\n"
             "L1：Patterned Narrative & Informational Readers\n"
             "L2：Early Independent Genre-Exposure Readers\n"
             "L3–L6：按实际为 fiction / non-fiction"],
            ["阅读字数", "取正文文本字数（不含题目）"],
            ["词汇难度", "写 CEFR 级别（如 Smart = Pre-A1）"],
            ["语法难度", "写时态、用中文表达（如：一般现在时态）"],
            ["词汇掌握", "master 的 4 个为主；高级别 4–6 个、可含词组——对应词汇掌握格子"],
            ["自然拼读（低级别）",
             "词组不首字母大写，用英文双引号 + 例词；例：consonant blend “fr” (friendship)、long “oy” (toy)、diphthong “ea+r” (bear)"],
            ["构词法（L5+ 取代拼读）",
             "如 suffix -ous (= having/full of a quality): nervous, famous, dangerous；例子给 1–2 个、不超一行"],
        ],
        widths=[1.9, 4.6],
    )

    # 4 Teacher's Guide
    h2(doc, "④ Teacher's Guide（教师指南 DOCX）")
    make_table(
        doc,
        ["维度", "标准 / 说明"],
        [
            ["定位", "给老师的授课指引"],
            ["模板", "严格套用官方母版（如 Spring Days 模板），Logo / 页眉页脚 / 版式 / 配色不动"],
            ["核心内容", "教学目标、目标词汇 / 句型、分步教学活动、课堂引导问题"],
            ["内容保真", "内容取自课程大纲；教学点 / 答案与其它三件套完全一致、不矛盾"],
        ],
        widths=[1.9, 4.6],
    )

    # 通用规范
    h2(doc, "⑤ 全交付物通用规范（100% 强制）")
    make_table(
        doc,
        ["规则", "标准 / 说明"],
        [
            ["命名规范", "统一格式：Level X_BookX_品类_标题.后缀；非法字符自动替换为下划线"],
            ["模板锁定", "严格套用原版母版，Logo / 页眉页脚 / 页边距 / 配色 / 装饰元素完全不动"],
            ["视觉统一", "全系列共用一套画风 + 固定 IP，跨页 / 跨册 / 跨交付物的人物、画风、配色完全一致"],
            ["内容保真", "文本优先取自课程大纲，无大纲不编造；跨交付物内容、答案一致无矛盾"],
            ["新人物规则", "故事出现大纲外新人物，须先确认形象设定再产出配图 / 内容"],
            ["生图 / 文本模型", "生图 gpt-image-2（imarouter 托管）；文本走 Claude；低随机、高提示词遵循度"],
            ["级别覆盖", "L0–L6 七个级别规则统一，差异仅在年龄 / 语言难度 / 词汇 / 句式复杂度"],
        ],
        widths=[1.6, 4.9],
    )

    dest = OUT / "绘本4件套_详细介绍.docx"
    doc.save(dest)
    return dest


if __name__ == "__main__":
    d1 = build_runsheet()
    print("SAVED", d1)
    d2 = build_suite_doc()
    print("SAVED", d2)
